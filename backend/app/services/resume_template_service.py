"""
Resume template service: parse resume text into structured sections,
then build a fully-formatted DOCX with distinct layouts per template.

8 Templates (inspired by popular LinkedIn / professional resume designs):
  1. Professional Classic  – single-column, centered name, horizontal rules (Calibri)
  2. Modern Tech           – two-column dark sidebar, blue accents (Arial)
  3. Creative Portfolio    – purple accents, decorative markers (Georgia)
  4. Academic CV           – formal "Curriculum Vitae" (Times New Roman)
  5. Executive             – navy+gold, premium spacing (Cambria)
  6. Minimalist Clean      – ultra-clean, monochrome, max whitespace (Calibri Light)
  7. Corporate Elegance    – teal header block, single-column body (Calibri)
  8. Fresh Graduate        – blue header bar, skills-first, compact (Arial)
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from typing import List, Optional

from sqlalchemy import text as sa_text
from sqlalchemy.orm import Session

_log = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────────────────────────
# Resume text parser
# ──────────────────────────────────────────────────────────────────────────────

_SECTION_KEYWORDS = [
    "summary", "profile", "objective", "about", "professional summary",
    "experience", "work experience", "professional experience", "employment",
    "education", "academic background",
    "skills", "technical skills", "core competencies", "competencies",
    "projects", "project experience",
    "certifications", "certificates", "licenses",
    "publications", "research",
    "awards", "honors", "achievements",
    "languages",
    "interests", "hobbies",
    "references", "volunteer", "volunteering",
    "activities", "extracurricular",
    "contact", "personal information",
]

# Chinese section titles (common resume headings)
_SECTION_KEYWORDS_ZH = [
    "个人简介", "简介", "摘要", "职业目标", "求职意向",
    "工作经历", "工作经验", "实习经历", "职业经历",
    "教育背景", "教育经历", "学历",
    "项目经历", "项目经验", "项目",
    "技能", "专业技能", "技术技能", "核心技能",
    "证书", "资格证书", "获奖", "荣誉", "奖项",
    "语言能力", "语言",
    "兴趣爱好", "兴趣",
    "自我评价", "联系方式", "个人信息",
    "发表论文", "研究成果", "学术成果",
]

_SECTION_RE = re.compile(
    r"^(?:" + "|".join(re.escape(k) for k in _SECTION_KEYWORDS) + r")\s*:?\s*$",
    re.IGNORECASE,
)

_SECTION_RE_ZH = re.compile(
    r"^(?:" + "|".join(re.escape(k) for k in _SECTION_KEYWORDS_ZH) + r")\s*:?\s*$",
)

# Words that indicate an ALL-CAPS line is a section heading (not a job title like "DATA SCIENTIST").
_SECTION_WORDS_HINT = frozenset(
    w.lower()
    for k in _SECTION_KEYWORDS
    for w in re.split(r"[\s/]+", k)
    if len(w) > 2
)

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_RE = re.compile(r"[\+]?[\d\s\-().]{7,15}")
_URL_RE = re.compile(r"(?:https?://|www\.)[\w./?=&%-]+", re.IGNORECASE)
# LinkedIn / GitHub etc. often appear without scheme
_SOCIAL_DOMAIN_RE = re.compile(
    r"(?:linkedin\.com|github\.com|gitlab\.com|behance\.net|notion\.so|medium\.com)(?:/[^\s]*)?",
    re.IGNORECASE,
)


@dataclass
class ResumeSection:
    title: str
    lines: List[str] = field(default_factory=list)


@dataclass
class ParsedResume:
    name: str = ""
    contact_lines: List[str] = field(default_factory=list)
    sections: List[ResumeSection] = field(default_factory=list)


def _is_section_header(line: str, known_name: Optional[str] = None) -> bool:
    stripped = line.strip().rstrip(":")
    if not stripped:
        return False
    if known_name and stripped.casefold() == known_name.strip().casefold():
        return False
    if _SECTION_RE.match(stripped):
        return True
    if _SECTION_RE_ZH.match(stripped):
        return True
    words = stripped.split()
    if 1 <= len(words) <= 6 and stripped == stripped.upper() and len(stripped) > 2:
        alpha_words = {w.lower() for w in words if w.isalpha()}
        if alpha_words & _SECTION_WORDS_HINT:
            return True
    if 1 <= len(words) <= 5 and all(w[0].isupper() for w in words if w.isalpha()):
        lower = stripped.lower().rstrip(":")
        for kw in _SECTION_KEYWORDS:
            if lower == kw or lower.startswith(kw + " "):
                return True
    return False


def _looks_like_contact(line: str) -> bool:
    s = line.strip()
    if _EMAIL_RE.search(s):
        return True
    if _URL_RE.search(s):
        return True
    if _SOCIAL_DOMAIN_RE.search(s):
        return True
    parts = [p.strip() for p in re.split(r"[|·•,]", s) if p.strip()]
    if len(parts) >= 2:
        has_contact = any(
            _EMAIL_RE.search(p)
            or _PHONE_RE.fullmatch(p.strip())
            or _URL_RE.search(p)
            or _SOCIAL_DOMAIN_RE.search(p)
            for p in parts
        )
        if has_contact:
            return True
    if _PHONE_RE.fullmatch(s):
        return True
    return False


def _is_plausible_resume_name(candidate: str) -> bool:
    """Heuristic: first line is a person name, not a headline or section."""
    c = candidate.strip()
    if not c:
        return False
    if len(c) > 60:
        return False
    if ":" in c:
        return False
    if "@" in c:
        return False
    if _EMAIL_RE.search(c) or _URL_RE.search(c) or _SOCIAL_DOMAIN_RE.search(c):
        return False
    # Looks like a sentence / objective, not a name
    if c.count(",") >= 2 or c.count("；") >= 1 or c.count(";") >= 2:
        return False
    lower = c.lower()
    if any(lower.startswith(kw + " ") for kw in _SECTION_KEYWORDS):
        return False
    if any(lower == kw for kw in _SECTION_KEYWORDS):
        return False
    # English regex header match (without circular title-case on name)
    if _SECTION_RE.match(c):
        return False
    if _SECTION_RE_ZH.match(c):
        return False
    # All-caps very long line: likely a headline, not a name
    if c == c.upper() and len(c) > 42:
        return False
    return True


def _normalize_resume_text(text: str) -> str:
    """Improve structure from mixed PDF/DOCX extraction: ZW chars, hyphen merges, spaces, long lines, blank runs."""
    if not text or not text.strip():
        return text
    text = re.sub(r"[\u200b-\u200d\ufeff]", "", text)
    raw_lines = text.splitlines()
    merged: List[str] = []
    i = 0
    while i < len(raw_lines):
        s = raw_lines[i].strip()
        s = re.sub(r"[ \t\xa0]+", " ", s) if s else s
        if not s:
            merged.append("")
            i += 1
            continue
        if s.endswith("-") and len(s) >= 2 and i + 1 < len(raw_lines):
            nxt = raw_lines[i + 1].strip()
            nxt = re.sub(r"[ \t\xa0]+", " ", nxt) if nxt else nxt
            if nxt and not nxt.startswith(("•", "-", "–", "▪", "►", "✦", "*", "▸")):
                merged.append(s[:-1].rstrip() + nxt)
                i += 2
                continue
        merged.append(s)
        i += 1

    out_lines: List[str] = []
    for line in merged:
        if not line.strip():
            out_lines.append("")
            continue
        s = line.strip()
        if len(s) > 400:
            parts = re.split(r"(?<=[.!?])\s+", s)
            for p in parts:
                if p.strip():
                    out_lines.append(p.strip())
        else:
            out_lines.append(s)
    collapsed: List[str] = []
    blank_run = 0
    for line in out_lines:
        if not line.strip():
            blank_run += 1
            if blank_run <= 2:
                collapsed.append("")
        else:
            blank_run = 0
            collapsed.append(line)
    return "\n".join(collapsed).strip()


def parse_resume(text: str) -> ParsedResume:
    """Split resume plain-text into name, contact lines, and titled sections."""
    text = _normalize_resume_text(text)
    lines = text.splitlines()
    result = ParsedResume()

    idx = 0
    while idx < len(lines) and not lines[idx].strip():
        idx += 1

    if idx < len(lines):
        candidate = lines[idx].strip()
        if candidate and not _is_section_header(candidate) and _is_plausible_resume_name(candidate):
            result.name = candidate
            idx += 1

    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue
        if _looks_like_contact(line):
            result.contact_lines.append(line)
            idx += 1
        elif _is_section_header(line, known_name=result.name or None):
            break
        elif len(result.contact_lines) == 0 and len(line) < 120:
            if "|" in line or "·" in line or "•" in line:
                result.contact_lines.append(line)
                idx += 1
            else:
                break
        else:
            break

    current_section: Optional[ResumeSection] = None
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if _is_section_header(stripped, known_name=result.name or None):
            if current_section is not None:
                result.sections.append(current_section)
            current_section = ResumeSection(title=stripped.rstrip(":").strip())
            idx += 1
            continue
        if current_section is None:
            if stripped:
                current_section = ResumeSection(title="Summary")
                current_section.lines.append(line)
        else:
            current_section.lines.append(line)
        idx += 1

    if current_section is not None:
        result.sections.append(current_section)

    for sec in result.sections:
        while sec.lines and not sec.lines[-1].strip():
            sec.lines.pop()

    return result


# ──────────────────────────────────────────────────────────────────────────────
# Shared helpers
# ──────────────────────────────────────────────────────────────────────────────

def _ensure_docx():
    try:
        from docx import Document  # noqa: F401
        return True
    except ImportError:
        raise RuntimeError("python-docx is required for template export")


def _set_cell_shading(cell, hex_color: str):
    from docx.oxml.ns import qn
    from lxml import etree
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shading = etree.SubElement(tc_pr, qn("w:shd"))
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins (in twips: 1pt = 20 twips)."""
    from docx.oxml.ns import qn
    from lxml import etree
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_mar = etree.SubElement(tc_pr, qn("w:tcMar"))
    for edge, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = etree.SubElement(tc_mar, qn(f"w:{edge}"))
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def _set_cell_vertical_alignment(cell, align="top"):
    """Set vertical alignment of table cell content."""
    from docx.oxml.ns import qn
    from lxml import etree
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:vAlign"))
    if existing is not None:
        tc_pr.remove(existing)
    v_align = etree.SubElement(tc_pr, qn("w:vAlign"))
    v_align.set(qn("w:val"), align)


def _ensure_cell_word_wrap(cell) -> None:
    """Allow text to wrap inside table cells (some generators set w:noWrap)."""
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for nw in list(tc_pr.findall(qn("w:noWrap"))):
        tc_pr.remove(nw)


def _set_paragraph_spacing(paragraph, before=0, after=0, line=None, line_multiple=None):
    """Body text: prefer ``line_multiple`` (e.g. 1.15) over fixed ``line`` (pt) for readability."""
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING

    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_multiple is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_multiple
    elif line is not None:
        pf.line_spacing = Pt(line)


def _set_line_spacing_multiple(paragraph, multiple: float = 1.15) -> None:
    from docx.enum.text import WD_LINE_SPACING

    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = multiple


def _add_page_number_footer(doc) -> None:
    """Centered PAGE field in footer (all templates)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    if p.text:
        p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char_b = OxmlElement("w:fldChar")
    fld_char_b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_e = OxmlElement("w:fldChar")
    fld_char_e.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_b)
    run._r.append(instr)
    run._r.append(fld_char_e)


def _add_header_name_centered(doc, name: str) -> None:
    """Optional header: applicant name (helps multi-page resumes in print preview)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    n = (name or "").strip()
    if not n:
        return
    section = doc.sections[0]
    hdr = section.header
    p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    if p.text:
        p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(n)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Calibri"


def _finalize_doc_header_footer(doc, parsed: ParsedResume) -> None:
    if parsed.name:
        _add_header_name_centered(doc, parsed.name)
    _add_page_number_footer(doc)


def _set_run_letter_spacing_twips(run, twips: int) -> None:
    """OOXML letter spacing (some python-docx versions ignore run.font.letter_spacing)."""
    from docx.oxml.ns import qn
    from lxml import etree

    r_pr = run._r.get_or_add_rPr()
    existing = r_pr.find(qn("w:spacing"))
    if existing is not None:
        r_pr.remove(existing)
    spacing = etree.SubElement(r_pr, qn("w:spacing"))
    spacing.set(qn("w:val"), str(twips))


def _emit_body_lines(
    add_paragraph,
    lines: List[str],
    *,
    font_name: str,
    normal_pt: float,
    sub_pt: float,
    normal_rgb,
    sub_rgb,
    bullet_left,
    bullet_hang,
    bullet_char: str = "•",
    sub_italic: bool = False,
    line_multiple: float = 1.15,
    body_left_indent=None,
    left_border_on_bullets: bool = False,
    left_border_hex: str = "0a2a4a",
    font_fallback: Optional[str] = None,
) -> None:
    """Shared body rendering: skip blank lines, unified bullets (•), 1.15 line spacing."""
    from docx.shared import Pt

    eff_fallback = font_fallback
    if eff_fallback is None:
        joined = "\n".join(lines)
        if _text_has_cjk(joined):
            eff_fallback = "Microsoft YaHei"

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        is_bullet = stripped.startswith(("•", "-", "–", "▪", "►", "✦", "*")) or bool(
            re.match(r"^\d{1,2}[\.\)]\s+\S", stripped)
        )
        is_sub = _is_sub_header(stripped) and not is_bullet
        p = add_paragraph()
        if is_bullet:
            text = re.sub(r"^\d{1,2}[\.\)]\s*", "", stripped)
            text = text.lstrip("•-–▪►✦* ").strip()
            p.paragraph_format.left_indent = bullet_left
            p.paragraph_format.first_line_indent = bullet_hang
            if left_border_on_bullets:
                _set_paragraph_left_border(p, left_border_hex, sz="8")
            run = p.add_run(f"{bullet_char}  {text}")
            run.font.size = Pt(normal_pt)
            run.font.color.rgb = normal_rgb
        elif is_sub:
            if body_left_indent is not None:
                p.paragraph_format.left_indent = body_left_indent
            run = p.add_run(stripped)
            run.font.size = Pt(sub_pt)
            run.font.bold = True
            run.font.italic = sub_italic
            run.font.color.rgb = sub_rgb
        else:
            if body_left_indent is not None:
                p.paragraph_format.left_indent = body_left_indent
            run = p.add_run(stripped)
            run.font.size = Pt(normal_pt)
            run.font.color.rgb = normal_rgb
        if eff_fallback:
            _set_run_font(run, font_name, eff_fallback)
        else:
            run.font.name = font_name
        _set_line_spacing_multiple(p, line_multiple)
        if is_bullet:
            _set_paragraph_spacing(p, before=1, after=2)
        elif is_sub:
            _set_paragraph_spacing(p, before=8, after=2)
        else:
            _set_paragraph_spacing(p, before=1, after=2)


def _set_paragraph_left_border(paragraph, color_hex: str = "0a2a4a", sz: str = "6") -> None:
    """Left vertical bar (academic-style emphasis); merges with existing pBdr if any."""
    from docx.oxml.ns import qn
    from lxml import etree

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = etree.SubElement(p_pr, qn("w:pBdr"))
    old_left = p_bdr.find(qn("w:left"))
    if old_left is not None:
        p_bdr.remove(old_left)
    left = etree.SubElement(p_bdr, qn("w:left"))
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)


def _set_paragraph_bottom_border(paragraph, color_hex, sz="4", space="1"):
    """Add a bottom border to a paragraph, removing any existing pBdr first."""
    from docx.oxml.ns import qn
    from lxml import etree
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color_hex)


def _add_horizontal_line(doc_or_cell, color_hex="444444", width_pt=0.5):
    p = doc_or_cell.add_paragraph()
    _set_paragraph_bottom_border(p, color_hex, sz=str(int(width_pt * 8)))
    _set_paragraph_spacing(p, before=2, after=4)
    return p


def _text_has_cjk(s: str) -> bool:
    """East Asian / CJK codepoints (incl. Japanese/Korean) for font fallback."""
    return bool(re.search(r"[\u3040-\u30ff\u3400-\u9fff\uac00-\ud7af]", s or ""))


def _set_run_font(run, font_name, fallback_name=None):
    """Set font name on a run with optional rFonts fallback for cross-platform + CJK."""
    run.font.name = font_name
    if fallback_name:
        from docx.oxml.ns import qn
        rPr = run._r.get_or_add_rPr()
        r_fonts = rPr.find(qn("w:rFonts"))
        if r_fonts is not None:
            r_fonts.set(qn("w:hAnsi"), font_name)
            r_fonts.set(qn("w:cs"), fallback_name)
            r_fonts.set(qn("w:eastAsia"), fallback_name)
        else:
            from lxml import etree
            r_fonts = etree.SubElement(rPr, qn("w:rFonts"))
            r_fonts.set(qn("w:ascii"), font_name)
            r_fonts.set(qn("w:hAnsi"), font_name)
            r_fonts.set(qn("w:cs"), fallback_name)
            r_fonts.set(qn("w:eastAsia"), fallback_name)


def _is_sub_header(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if re.match(r"^.{3,60}\s*[|–—-]\s*.{3,60}$", stripped):
        return True
    if re.match(r"^.{3,50}\s*\(\s*\d{4}", stripped):
        return True
    if stripped.endswith(":") and len(stripped) < 60 and not _is_section_header(stripped):
        return True
    return False


def _split_contact_parts(contact_lines: List[str]) -> List[str]:
    """Flatten contact lines separated by | . into individual items."""
    parts = []
    for cl in contact_lines:
        for p in re.split(r"[|·•]", cl):
            p = p.strip()
            if p:
                parts.append(p)
    return parts


_SIDEBAR_SECTION_NAMES = {
    "skills", "technical skills", "core competencies", "competencies",
    "languages", "certifications", "certificates", "licenses",
    "interests", "hobbies", "contact", "personal information",
    "awards", "honors", "achievements",
}


_SIDEBAR_SECTION_NAMES_ZH = frozenset(
    {
        "技能",
        "专业技能",
        "技术技能",
        "语言能力",
        "证书",
        "资格证书",
        "兴趣爱好",
        "兴趣",
        "获奖",
        "荣誉",
        "奖项",
        "联系方式",
        "个人信息",
    }
)


def _partition_sections(parsed: ParsedResume):
    """Split sections into sidebar vs main content."""
    sidebar, main = [], []
    for sec in parsed.sections:
        t = sec.title.strip()
        if t.lower() in _SIDEBAR_SECTION_NAMES or t in _SIDEBAR_SECTION_NAMES_ZH:
            sidebar.append(sec)
        else:
            main.append(sec)
    return sidebar, main


def _remove_table_borders(table):
    """Remove all borders from a table."""
    from docx.oxml.ns import qn
    from lxml import etree
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn("w:tblPr"))
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = etree.SubElement(tbl_pr, qn("w:tblBorders"))
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = etree.SubElement(borders, qn(f"w:{edge}"))
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")


def _is_skills_section(title: str) -> bool:
    t = title.strip().lower()
    if t in {"skills", "technical skills", "core competencies", "competencies"}:
        return True
    return title.strip() in {"技能", "专业技能", "技术技能", "核心技能"}


def _format_skills_inline(lines: List[str]) -> str:
    """Convert skills lines into readable text; preserve 'Label: a, b' as structured segments."""
    structured: List[str] = []
    flat: List[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        stripped = stripped.lstrip("•-–▪►✦*▸ ").strip()
        if not stripped:
            continue
        if re.match(r"^[^:]{1,48}:\s*.{2,}", stripped) and "http" not in stripped.lower():
            structured.append(stripped)
            continue
        for part in re.split(r"[,;]", stripped):
            part = part.strip()
            if part:
                flat.append(part)
    parts_out: List[str] = []
    if structured:
        parts_out.append(" ".join(structured))
    if flat:
        parts_out.append(", ".join(flat))
    return " · ".join(parts_out) if parts_out else ""


def _emit_fresh_graduate_skills_section(
    add_paragraph,
    lines: List[str],
    *,
    font_name: str,
    normal_pt: float,
    gray_rgb,
    dark_rgb,
) -> None:
    """Skills block with bold category labels (Programming: …) plus flat comma-separated items."""
    from docx.shared import Pt

    structured: List[str] = []
    flat: List[str] = []
    for line in lines:
        stripped = line.strip().lstrip("•-–▪►✦*▸ ").strip()
        if not stripped:
            continue
        if re.match(r"^[^:]{1,48}:\s*.{2,}", stripped) and "http" not in stripped.lower():
            structured.append(stripped)
            continue
        for part in re.split(r"[,;]", stripped):
            part = part.strip()
            if part:
                flat.append(part)

    for sl in structured:
        if ":" not in sl:
            continue
        label, rest = sl.split(":", 1)
        p = add_paragraph()
        r1 = p.add_run(label.strip() + ":")
        r1.font.bold = True
        r1.font.size = Pt(normal_pt)
        r1.font.color.rgb = dark_rgb
        r1.font.name = font_name
        r2 = p.add_run(rest)
        r2.font.size = Pt(normal_pt)
        r2.font.color.rgb = gray_rgb
        r2.font.name = font_name
        _set_line_spacing_multiple(p, 1.15)
        _set_paragraph_spacing(p, before=3, after=4)

    if flat:
        p = add_paragraph()
        run = p.add_run(", ".join(flat))
        run.font.size = Pt(normal_pt)
        run.font.color.rgb = gray_rgb
        run.font.name = font_name
        _set_line_spacing_multiple(p, 1.15)
        _set_paragraph_spacing(p, before=3, after=3)


def _get_first_paragraph(cell):
    """Get first paragraph of a cell, reusing the existing empty one if possible."""
    if cell.paragraphs and not cell.paragraphs[0].text:
        return cell.paragraphs[0]
    return cell.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# Template 1: Professional Classic
# Single-column, centered name, horizontal rules. ATS-friendly.
# Font: Calibri (universal). Inspired by Resume.io "London" / Zety "Traditional".
# ══════════════════════════════════════════════════════════════════════════════

def _build_professional_classic(parsed: ParsedResume) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x1a, 0x1a, 0x2e)
    DARK = RGBColor(0x33, 0x33, 0x33)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(2)

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    if parsed.name:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(parsed.name.upper())
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY
        run.font.bold = True
        run.font.name = "Calibri"
        _set_run_letter_spacing_twips(run, 40)
        _set_paragraph_spacing(h, before=0, after=4)

    if parsed.contact_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = "  |  ".join(_split_contact_parts(parsed.contact_lines))
        run = p.add_run(contact_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = "Calibri"
        _set_run_letter_spacing_twips(run, 6)
        _set_paragraph_spacing(p, before=0, after=6)

    _add_horizontal_line(doc, "1a1a2e", 1)

    for section in parsed.sections:
        p = doc.add_paragraph()
        run = p.add_run(section.title.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
        _set_run_letter_spacing_twips(run, 30)
        _set_paragraph_spacing(p, before=10, after=6)
        _add_horizontal_line(doc, "cccccc", 0.5)

        _emit_body_lines(
            doc.add_paragraph,
            section.lines,
            font_name="Calibri",
            normal_pt=10.5,
            sub_pt=11,
            normal_rgb=DARK,
            sub_rgb=RGBColor(0x2a, 0x2a, 0x2a),
            bullet_left=Inches(0.3),
            bullet_hang=Inches(-0.15),
        )

    _finalize_doc_header_footer(doc, parsed)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# Template 2: Modern Tech (two-column, dark sidebar)
# Font: Arial (universal fallback for Segoe UI). Inspired by Resume.io "Dublin".
# ══════════════════════════════════════════════════════════════════════════════

def _build_modern_tech(parsed: ParsedResume) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches

    SIDEBAR_BG = "0f172a"
    ACCENT = RGBColor(0x38, 0xbd, 0xf8)
    WHITE = RGBColor(0xff, 0xff, 0xff)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)

    sec = doc.sections[0]
    sec.top_margin = Cm(0)
    sec.bottom_margin = Cm(1)
    sec.left_margin = Cm(0)
    sec.right_margin = Cm(0)

    sidebar_sections, main_sections = _partition_sections(parsed)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    _remove_table_borders(table)

    page_width = sec.page_width
    sidebar_w = int(page_width * 0.32)
    main_w = page_width - sidebar_w
    table.columns[0].width = sidebar_w
    table.columns[1].width = main_w

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    left_cell.width = sidebar_w
    right_cell.width = main_w
    _set_cell_shading(left_cell, SIDEBAR_BG)
    _set_cell_margins(left_cell, top=300, bottom=300, left=170, right=100)
    _set_cell_vertical_alignment(left_cell, "top")
    _set_cell_margins(right_cell, top=300, bottom=200, left=200, right=200)
    _set_cell_vertical_alignment(right_cell, "top")
    _ensure_cell_word_wrap(left_cell)
    _ensure_cell_word_wrap(right_cell)

    for p in left_cell.paragraphs:
        p.text = ""

    if parsed.name:
        p = _get_first_paragraph(left_cell)
        run = p.add_run(parsed.name)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.name = "Arial"
        _set_paragraph_spacing(p, before=6, after=4)
        p.paragraph_format.left_indent = Inches(0.15)

    if parsed.contact_lines:
        for part in _split_contact_parts(parsed.contact_lines):
            p = left_cell.add_paragraph()
            run = p.add_run(part)
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
            run.font.name = "Arial"
            _set_paragraph_spacing(p, before=1, after=1)
            p.paragraph_format.left_indent = Inches(0.15)

    for ssec in sidebar_sections:
        p = left_cell.add_paragraph()
        _set_paragraph_spacing(p, before=12, after=2)
        p.paragraph_format.left_indent = Inches(0.15)
        _set_paragraph_bottom_border(p, "38bdf8", sz="3")

        p = left_cell.add_paragraph()
        run = p.add_run(ssec.title.upper())
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = ACCENT
        run.font.name = "Arial"
        _set_run_letter_spacing_twips(run, 20)
        _set_paragraph_spacing(p, before=1, after=4)
        p.paragraph_format.left_indent = Inches(0.15)

        _emit_body_lines(
            left_cell.add_paragraph,
            ssec.lines,
            font_name="Arial",
            normal_pt=9,
            sub_pt=9.5,
            normal_rgb=RGBColor(0xcb, 0xd5, 0xe1),
            sub_rgb=RGBColor(0xee, 0xf2, 0xf6),
            bullet_left=Inches(0.2),
            bullet_hang=Inches(-0.12),
        )

    p = left_cell.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=10)

    _set_cell_shading(right_cell, "ffffff")
    for p in right_cell.paragraphs:
        p.text = ""

    first_main = True
    for msec in main_sections:
        p = _get_first_paragraph(right_cell) if first_main else right_cell.add_paragraph()
        first_main = False
        run = p.add_run(msec.title.upper())
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        run.font.name = "Arial"
        _set_run_letter_spacing_twips(run, 30)
        _set_paragraph_spacing(p, before=10, after=2)
        p.paragraph_format.left_indent = Inches(0.1)

        p = right_cell.add_paragraph()
        _set_paragraph_bottom_border(p, "38bdf8", sz="4")
        _set_paragraph_spacing(p, before=0, after=6)
        p.paragraph_format.left_indent = Inches(0.1)

        def _add_paragraph_main():
            p2 = right_cell.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.1)
            return p2

        _emit_body_lines(
            right_cell.add_paragraph,
            msec.lines,
            font_name="Arial",
            normal_pt=10,
            sub_pt=10.5,
            normal_rgb=RGBColor(0x44, 0x44, 0x44),
            sub_rgb=RGBColor(0x1e, 0x29, 0x3b),
            bullet_left=Inches(0.35),
            bullet_hang=Inches(-0.15),
            body_left_indent=Inches(0.1),
        )

    _finalize_doc_header_footer(doc, parsed)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# Template 3: Creative Portfolio
# Font: Georgia (universal serif). Inspired by Novoresume "Creative".
# ══════════════════════════════════════════════════════════════════════════════

def _build_creative_portfolio(parsed: ParsedResume) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches

    PURPLE = RGBColor(0x7c, 0x3a, 0xed)
    DARK = RGBColor(0x2d, 0x2d, 0x2d)
    GRAY = RGBColor(0x55, 0x55, 0x55)
    LIGHT_PURPLE = "c4b5fd"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    if parsed.name:
        h = doc.add_paragraph()
        name = parsed.name.strip()
        first = name[:1]
        rest = name[1:]
        use_drop_cap = (
            len(rest) > 0
            and first.isascii()
            and first.isalpha()
            and len(name) <= 80
        )
        if use_drop_cap:
            run1 = h.add_run(first)
            run1.font.size = Pt(36)
            run1.font.bold = True
            run1.font.color.rgb = PURPLE
            run1.font.name = "Georgia"
            run2 = h.add_run(rest)
            run2.font.size = Pt(28)
            run2.font.color.rgb = DARK
            run2.font.name = "Georgia"
        else:
            run = h.add_run(name)
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = PURPLE
            run.font.name = "Georgia"
        _set_paragraph_spacing(h, before=0, after=2)

    # Purple accent bar
    p = doc.add_paragraph()
    _set_paragraph_bottom_border(p, "7c3aed", sz="12")
    _set_paragraph_spacing(p, before=0, after=4)

    if parsed.contact_lines:
        p = doc.add_paragraph()
        run = p.add_run("  |  ".join(_split_contact_parts(parsed.contact_lines)))
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.font.name = "Georgia"
        run.font.italic = True
        _set_paragraph_spacing(p, before=0, after=10)

    for section in parsed.sections:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=14, after=2)
        run = p.add_run(section.title)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = PURPLE
        run.font.name = "Georgia"

        # Light purple underline
        _set_paragraph_bottom_border(p, LIGHT_PURPLE, sz="4")

        _emit_body_lines(
            doc.add_paragraph,
            section.lines,
            font_name="Georgia",
            normal_pt=10.5,
            sub_pt=11,
            normal_rgb=DARK,
            sub_rgb=RGBColor(0x4a, 0x4a, 0x4a),
            bullet_left=Inches(0.4),
            bullet_hang=Inches(-0.15),
            body_left_indent=Inches(0.2),
            sub_italic=True,
        )

    _finalize_doc_header_footer(doc, parsed)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# Template 4: Academic CV
# Font: Times New Roman (universal). Inspired by classic CV format.
# ══════════════════════════════════════════════════════════════════════════════

def _build_academic_research(parsed: ParsedResume) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x0a, 0x2a, 0x4a)
    DARK = RGBColor(0x1a, 0x1a, 0x1a)
    GRAY = RGBColor(0x44, 0x44, 0x44)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(2)

    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("CURRICULUM VITAE")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.name = "Times New Roman"
    _set_run_letter_spacing_twips(run, 60)
    _set_paragraph_spacing(h, before=0, after=6)

    _add_horizontal_line(doc, "0a2a4a", 1.5)

    if parsed.name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(parsed.name)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = DARK
        run.font.name = "Times New Roman"
        _set_paragraph_spacing(p, before=4, after=4)

    if parsed.contact_lines:
        for cl in parsed.contact_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cl)
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY
            run.font.name = "Times New Roman"
            _set_paragraph_spacing(p, before=0, after=2)
        _add_horizontal_line(doc, "0a2a4a", 0.5)

    for section in parsed.sections:
        p = doc.add_paragraph()
        run = p.add_run(section.title.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run.font.name = "Times New Roman"
        _set_run_letter_spacing_twips(run, 20)
        _set_paragraph_spacing(p, before=14, after=2)
        _set_paragraph_bottom_border(p, "0a2a4a", sz="4")

        _emit_body_lines(
            doc.add_paragraph,
            section.lines,
            font_name="Times New Roman",
            normal_pt=11,
            sub_pt=11,
            normal_rgb=DARK,
            sub_rgb=DARK,
            bullet_left=Inches(0.5),
            bullet_hang=Inches(-0.2),
            body_left_indent=Inches(0.3),
            left_border_on_bullets=True,
            left_border_hex="0a2a4a",
        )

    _finalize_doc_header_footer(doc, parsed)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# Template 5: Executive (navy + gold, premium spacing)
# Font: Cambria (universal serif, replaces Garamond). Inspired by Resume.io "Sterling".
# ══════════════════════════════════════════════════════════════════════════════

def _build_executive(parsed: ParsedResume) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x1b, 0x2a, 0x4a)
    GOLD = RGBColor(0xbf, 0x94, 0x3e)
    DARK = RGBColor(0x22, 0x22, 0x22)
    GRAY = RGBColor(0x66, 0x66, 0x66)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Cambria"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(3)

    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(3)

    # Gold top border
    p = doc.add_paragraph()
    _set_paragraph_bottom_border(p, "bf943e", sz="12")
    _set_paragraph_spacing(p, before=0, after=10)

    if parsed.name:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(parsed.name.upper())
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = NAVY
        _set_run_font(run, "Cambria", "Georgia")
        _set_run_letter_spacing_twips(run, 80)
        _set_paragraph_spacing(h, before=0, after=6)

    if parsed.contact_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = _split_contact_parts(parsed.contact_lines)
        run = p.add_run("  |  ".join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        _set_run_font(run, "Cambria", "Georgia")
        _set_paragraph_spacing(p, before=0, after=4)

    # Gold bottom border
    p = doc.add_paragraph()
    _set_paragraph_bottom_border(p, "bf943e", sz="12")
    _set_paragraph_spacing(p, before=4, after=12)

    for section in parsed.sections:
        p = doc.add_paragraph()
        run = p.add_run(section.title.upper())
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = NAVY
        _set_run_font(run, "Cambria", "Georgia")
        _set_run_letter_spacing_twips(run, 40)
        _set_paragraph_spacing(p, before=12, after=2)
        _set_paragraph_bottom_border(p, "bf943e", sz="6", space="2")
        p_gray = doc.add_paragraph()
        _set_paragraph_bottom_border(p_gray, "888888", sz="2")
        _set_paragraph_spacing(p_gray, before=0, after=6)

        _emit_body_lines(
            doc.add_paragraph,
            section.lines,
            font_name="Cambria",
            font_fallback="Georgia",
            normal_pt=11,
            sub_pt=11.5,
            normal_rgb=DARK,
            sub_rgb=NAVY,
            bullet_left=Inches(0.4),
            bullet_hang=Inches(-0.2),
            body_left_indent=Inches(0.1),
        )

    _finalize_doc_header_footer(doc, parsed)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# Template 6: Minimalist Clean
# Font: Calibri Light (universal, replaces Helvetica). #1 most downloaded style.
# Inspired by Canva "Black White Minimalist" / Resume.io "Essential".
# ══════════════════════════════════════════════════════════════════════════════

def _build_minimalist_clean(parsed: ParsedResume) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches

    BLACK = RGBColor(0x11, 0x11, 0x11)
    MID = RGBColor(0x55, 0x55, 0x55)
    LIGHT = RGBColor(0x99, 0x99, 0x99)
    FONT = "Calibri Light"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style.font.color.rgb = MID
    style.paragraph_format.space_after = Pt(2)

    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(3)

    if parsed.name:
        h = doc.add_paragraph()
        run = h.add_run(parsed.name)
        run.font.size = Pt(28)
        run.font.color.rgb = BLACK
        run.font.name = FONT
        _set_paragraph_spacing(h, before=0, after=4)

    if parsed.contact_lines:
        p = doc.add_paragraph()
        parts = _split_contact_parts(parsed.contact_lines)
        run = p.add_run("    ".join(parts))
        run.font.size = Pt(8.5)
        run.font.color.rgb = LIGHT
        run.font.name = FONT
        _set_paragraph_spacing(p, before=0, after=9)

    for section in parsed.sections:
        p = doc.add_paragraph()
        run = p.add_run(section.title.lower())
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = BLACK
        run.font.name = FONT
        _set_run_letter_spacing_twips(run, 40)
        _set_paragraph_spacing(p, before=16, after=6)
        _add_horizontal_line(doc, "dddddd", 0.25)

        _emit_body_lines(
            doc.add_paragraph,
            section.lines,
            font_name=FONT,
            normal_pt=10,
            sub_pt=10.5,
            normal_rgb=MID,
            sub_rgb=BLACK,
            bullet_left=Inches(0.2),
            bullet_hang=Inches(-0.12),
        )

    _finalize_doc_header_footer(doc, parsed)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# Template 7: Corporate Elegance (single-column with teal header block)
# Restructured from two-column to single-column for better ATS compatibility
# and visual differentiation from Modern Tech.
# Font: Calibri (universal). Inspired by Resume.io "Corporate" / Zety "Elegant".
# ══════════════════════════════════════════════════════════════════════════════

def _build_corporate_elegance(parsed: ParsedResume) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    TEAL_BG = "134e4a"
    TEAL = RGBColor(0x14, 0xb8, 0xa6)
    DARK = RGBColor(0x1f, 0x2a, 0x37)
    GRAY = RGBColor(0x4b, 0x55, 0x63)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK

    sec = doc.sections[0]
    sec.top_margin = Cm(0)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # Teal header block using full-width table
    header_tbl = doc.add_table(rows=1, cols=1)
    header_tbl.autofit = True
    _remove_table_borders(header_tbl)
    header_cell = header_tbl.cell(0, 0)
    _set_cell_shading(header_cell, TEAL_BG)
    _set_cell_margins(header_cell, top=350, bottom=350, left=300, right=300)

    for p in header_cell.paragraphs:
        p.text = ""

    if parsed.name:
        p = _get_first_paragraph(header_cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(parsed.name.upper())
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run.font.name = "Calibri"
        _set_run_letter_spacing_twips(run, 60)
        _set_paragraph_spacing(p, before=4, after=4)

    if parsed.contact_lines:
        p = header_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = _split_contact_parts(parsed.contact_lines)
        run = p.add_run("  |  ".join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xa7, 0xf3, 0xd0)
        run.font.name = "Calibri"
        _set_paragraph_spacing(p, before=0, after=4)

    # Teal accent line after header
    p = doc.add_paragraph()
    _set_paragraph_bottom_border(p, "14b8a6", sz="8")
    _set_paragraph_spacing(p, before=0, after=8)

    for section in parsed.sections:
        p = doc.add_paragraph()
        run = p.add_run(section.title.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x13, 0x4e, 0x4a)
        run.font.name = "Calibri"
        _set_run_letter_spacing_twips(run, 30)
        _set_paragraph_spacing(p, before=12, after=2)
        _set_paragraph_bottom_border(p, "14b8a6", sz="3")

        _emit_body_lines(
            doc.add_paragraph,
            section.lines,
            font_name="Calibri",
            normal_pt=10.5,
            sub_pt=11,
            normal_rgb=GRAY,
            sub_rgb=DARK,
            bullet_left=Inches(0.35),
            bullet_hang=Inches(-0.15),
        )

    _finalize_doc_header_footer(doc, parsed)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# Template 8: Fresh Graduate (compact, friendly, skills-first)
# Font: Arial (universal). Inspired by Novoresume "Functional Modern".
# Skills sections are formatted as comma-separated inline text.
# ══════════════════════════════════════════════════════════════════════════════

def _build_fresh_graduate(parsed: ParsedResume) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    BLUE = RGBColor(0x25, 0x63, 0xeb)
    DARK = RGBColor(0x1e, 0x29, 0x3b)
    GRAY = RGBColor(0x4b, 0x55, 0x63)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    sec = doc.sections[0]
    sec.top_margin = Cm(0)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    # Blue header bar
    header_tbl = doc.add_table(rows=1, cols=1)
    header_tbl.autofit = True
    _remove_table_borders(header_tbl)
    header_cell = header_tbl.cell(0, 0)
    _set_cell_shading(header_cell, "2563eb")
    _set_cell_margins(header_cell, top=250, bottom=250, left=200, right=200)

    for p in header_cell.paragraphs:
        p.text = ""

    if parsed.name:
        p = _get_first_paragraph(header_cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(parsed.name.upper())
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run.font.name = "Arial"
        _set_run_letter_spacing_twips(run, 60)
        _set_paragraph_spacing(p, before=4, after=4)

    if parsed.contact_lines:
        p = header_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = _split_contact_parts(parsed.contact_lines)
        run = p.add_run("  |  ".join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xdb, 0xea, 0xfe)
        run.font.name = "Arial"
        _set_paragraph_spacing(p, before=0, after=4)

    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=4, after=4)

    # Reorder: put skills/education before experience for fresh grads
    skills_first = []
    experience_later = []
    for s in parsed.sections:
        title_lower = s.title.lower()
        zh = s.title.strip()
        if any(k in title_lower for k in ("skill", "education", "certif", "language", "award")) or any(
            x in zh for x in ("技能", "教育", "证书", "语言", "获奖", "荣誉")
        ):
            skills_first.append(s)
        else:
            experience_later.append(s)
    ordered_sections = skills_first + experience_later

    for section in ordered_sections:
        p = doc.add_paragraph()
        run = p.add_run(section.title.upper())
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = DARK
        run.font.name = "Arial"
        _set_run_letter_spacing_twips(run, 20)
        _set_paragraph_spacing(p, before=10, after=2)
        _set_paragraph_bottom_border(p, "93c5fd", sz="4")

        if _is_skills_section(section.title):
            _emit_fresh_graduate_skills_section(
                doc.add_paragraph,
                section.lines,
                font_name="Arial",
                normal_pt=10,
                gray_rgb=GRAY,
                dark_rgb=DARK,
            )
            continue

        _emit_body_lines(
            doc.add_paragraph,
            section.lines,
            font_name="Arial",
            normal_pt=10,
            sub_pt=10.5,
            normal_rgb=GRAY,
            sub_rgb=DARK,
            bullet_left=Inches(0.3),
            bullet_hang=Inches(-0.15),
        )

    _finalize_doc_header_footer(doc, parsed)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# Public API
# ══════════════════════════════════════════════════════════════════════════════

_TEMPLATE_BUILDERS = {
    "professional_classic": _build_professional_classic,
    "modern_tech": _build_modern_tech,
    "creative_portfolio": _build_creative_portfolio,
    "academic_research": _build_academic_research,
    "executive": _build_executive,
    "minimalist_clean": _build_minimalist_clean,
    "corporate_elegance": _build_corporate_elegance,
    "fresh_graduate": _build_fresh_graduate,
}


def resolve_template_builder_key(template_id: str, db: Session) -> str:
    """Public helper for preview / analytics — maps DB template id to builder key."""
    return _resolve_template_key(template_id, db)


def _resolve_template_key(template_id: str, db: Session) -> str:
    """Map template_id to a builder key."""
    clean = template_id.strip().lstrip("_")
    for key in _TEMPLATE_BUILDERS:
        if clean == key or clean.replace(".docx", "") == key:
            return key

    row = db.execute(
        sa_text("SELECT template_file FROM resume_templates WHERE template_id = :tid AND is_active = TRUE LIMIT 1"),
        {"tid": template_id},
    ).mappings().first()
    if row:
        fname = (row.get("template_file") or "").replace(".docx", "").strip()
        if fname in _TEMPLATE_BUILDERS:
            return fname

    return "professional_classic"


def apply_template(
    db: Session,
    review_id: str,
    template_id: str,
    resume_content: str,
    template_file: Optional[str] = None,
) -> bytes:
    """
    Parse resume_content into structured sections, then build a fully-formatted
    DOCX using the template identified by template_id.

    Layout is generated in code (``Document()`` + builders); disk ``.docx`` seeds are
    not merged. Shared body rendering uses :func:`_emit_body_lines` (1.15 line spacing,
    unified bullets, optional left bar for academic). Header shows the parsed name when
    available; footer includes page numbers. Word custom styles can be introduced
    gradually to reduce per-run formatting.
    """
    _ensure_docx()

    key = _resolve_template_key(template_id, db)

    if template_file:
        fname = template_file.replace(".docx", "").strip()
        if fname in _TEMPLATE_BUILDERS:
            key = fname

    _log.info("apply_template: using builder '%s' for template_id='%s'", key, template_id)

    parsed = parse_resume(resume_content or "")
    try:
        from backend.app.services.resume_structured import enhance_parsed_for_export

        parsed = enhance_parsed_for_export(parsed)
    except Exception as ex:
        _log.warning("enhance_parsed_for_export skipped: %s", ex)
    _log.info(
        "apply_template: parsed name='%s', %d contact lines, %d sections",
        parsed.name, len(parsed.contact_lines), len(parsed.sections),
    )

    builder = _TEMPLATE_BUILDERS.get(key, _build_professional_classic)
    return builder(parsed)
