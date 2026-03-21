import logging
from sqlalchemy import text
from fastapi import Depends, FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from fastapi.responses import JSONResponse
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import Response

from backend.app.db.deps import get_db
from backend.app.db.session import SessionLocal
from backend.app.middleware.audit_middleware import AuditMiddleware
from backend.app.middleware.rate_limit_middleware import RateLimitMiddleware
from backend.app.rate_limit import _parse_bool_env
from backend.app.security import require_production_secret, _is_production

# SkillSight schemas (Week1 Day3)
# NOTE: repo-root/ is expected to be on PYTHONPATH when running uvicorn from repo root.
try:
    from schemas.skillsight_models import Skill, Role, EvidencePointer, AuditLog, ConsentRecord  # noqa: F401
except ImportError:
    pass  # Schemas optional for demo

app = FastAPI(title="SkillSight API", version="0.1.0")

logger = logging.getLogger(__name__)
_SCHEMA_HEALTH: dict = {"ok": True, "missing": [], "checked_at": None}


def _run_schema_health_check() -> dict:
    required = {
        "assessment_sessions": ["session_id", "user_id", "assessment_type", "status", "config"],
        "assessment_attempts": ["attempt_id", "session_id", "evaluation", "score"],
        "skill_assessments": ["assessment_id", "doc_id", "skill_id", "decision", "decision_meta"],
        "skill_proficiency": ["prof_id", "doc_id", "skill_id", "level", "label"],
        "consents": ["doc_id", "user_id", "status"],
    }
    missing = []
    try:
        with SessionLocal() as db:
            rows = db.execute(text("""
                SELECT table_name, column_name
                FROM information_schema.columns
                WHERE table_schema = 'public'
            """)).mappings().all()
        by_table = {}
        for row in rows:
            by_table.setdefault(str(row["table_name"]), set()).add(str(row["column_name"]))
        for table_name, cols in required.items():
            if table_name not in by_table:
                missing.append(f"missing_table:{table_name}")
                continue
            for col in cols:
                if col not in by_table[table_name]:
                    missing.append(f"missing_column:{table_name}.{col}")
    except Exception as e:
        missing.append(f"schema_check_error:{type(e).__name__}:{e}")

    return {"ok": len(missing) == 0, "missing": missing, "checked_at": None}


def _seed_roles_and_skills(db):
    """Seed roles + skills from JSON data files if DB tables are empty."""
    import json, pathlib, uuid
    from datetime import datetime, timezone
    from sqlalchemy import text

    base = pathlib.Path(__file__).resolve().parent.parent / "data"
    now = datetime.now(timezone.utc)

    # --- seed skills ---
    skills_file = (base / "seeds" / "skills_comprehensive.json") if (base / "seeds" / "skills_comprehensive.json").exists() else (base / "skills.json")
    if skills_file.exists():
        try:
            count = db.execute(text("SELECT count(*) FROM skills")).scalar() or 0
        except Exception as exc:
            logger.debug("skills table not ready: %s", exc)
            count = -1
        if count == 0:
            skills = json.loads(skills_file.read_text(encoding="utf-8"))
            for s in skills:
                sid = s.get("skill_id", "").strip()
                if not sid:
                    continue
                db.execute(
                    text("""
                        INSERT INTO skills (skill_id, canonical_name, definition,
                                            evidence_rules, level_rubric_json, version, source, created_at)
                        VALUES (:skill_id, :canonical_name, :definition,
                                :evidence_rules, :level_rubric_json, :version, :source, :created_at)
                        ON CONFLICT (skill_id) DO NOTHING
                    """),
                    {
                        "skill_id": sid,
                        "canonical_name": s.get("canonical_name", ""),
                        "definition": s.get("definition", ""),
                        "evidence_rules": s.get("evidence_rules", ""),
                        "level_rubric_json": json.dumps(s.get("level_rubric", {}), ensure_ascii=False),
                        "version": s.get("version", "v1"),
                        "source": s.get("source", "HKU"),
                        "created_at": now,
                    },
                )
                for alias in s.get("aliases", []):
                    try:
                        db.execute(
                            text("""
                                INSERT INTO skill_aliases (alias_id, skill_id, alias, source, created_at)
                                VALUES ((:aid)::uuid, :skill_id, :alias, :source, :created_at)
                                ON CONFLICT (skill_id, alias) DO NOTHING
                            """),
                            {
                                "aid": str(uuid.uuid4()),
                                "skill_id": sid,
                                "alias": alias,
                                "source": s.get("source", "HKU"),
                                "created_at": now,
                            },
                        )
                    except Exception as ex:
                        logger.debug("skill_aliases seed skip skill=%s alias=%s: %s", sid, alias, ex)
            logger.info("Seeded %d skills from %s", len(skills), skills_file.name)

    # --- seed roles ---
    roles_file = base / "roles.json"
    if roles_file.exists():
        try:
            count = db.execute(text("SELECT count(*) FROM roles")).scalar() or 0
        except Exception as exc:
            logger.debug("roles table not ready: %s", exc)
            count = -1
        if count == 0:
            roles = json.loads(roles_file.read_text(encoding="utf-8"))
            for r in roles:
                rid = r.get("role_id", "").strip()
                if not rid:
                    continue
                db.execute(
                    text("""
                        INSERT INTO roles (role_id, role_title, description, created_at, updated_at)
                        VALUES (:role_id, :role_title, :description, :created_at, :updated_at)
                        ON CONFLICT (role_id) DO NOTHING
                    """),
                    {
                        "role_id": rid,
                        "role_title": r.get("role_title", ""),
                        "description": r.get("description", ""),
                        "created_at": now,
                        "updated_at": now,
                    },
                )
                for sr in r.get("skills_required", []):
                    skill_id = sr.get("skill_id", "").strip()
                    if not skill_id:
                        continue
                    db.execute(
                        text("""
                            INSERT INTO role_skill_requirements
                                   (req_id, role_id, skill_id, target_level, required, weight, created_at)
                            VALUES ((:req_id)::uuid, :role_id, :skill_id, :target_level, :required, :weight, :created_at)
                            ON CONFLICT (role_id, skill_id) DO UPDATE
                                SET target_level=EXCLUDED.target_level, required=EXCLUDED.required, weight=EXCLUDED.weight
                        """),
                        {
                            "req_id": str(uuid.uuid4()),
                            "role_id": rid,
                            "skill_id": skill_id,
                            "target_level": int(sr.get("target_level", 0)),
                            "required": bool(sr.get("required", True)),
                            "weight": float(sr.get("weight", 1.0)),
                            "created_at": now,
                        },
                    )
            logger.info("Seeded %d roles from roles.json", len(roles))


def _seed_resume_templates(db):
    """Seed resume templates from JSON file, inserting any missing templates."""
    import json, pathlib
    from sqlalchemy import text

    seed_file = pathlib.Path(__file__).resolve().parent.parent / "data" / "seeds" / "resume_templates.json"
    if not seed_file.exists():
        return

    try:
        count = db.execute(text("SELECT count(*) FROM resume_templates")).scalar() or 0
    except Exception:
        return

    templates = json.loads(seed_file.read_text(encoding="utf-8"))
    if count >= len(templates):
        return

    existing_files = set()
    try:
        rows = db.execute(text("SELECT template_file FROM resume_templates")).mappings().all()
        existing_files = {r["template_file"] for r in rows if r.get("template_file")}
    except Exception:
        pass

    inserted = 0
    for t in templates:
        tfile = t.get("template_file", "")
        if tfile in existing_files:
            continue
        try:
            db.execute(
                text("""
                    INSERT INTO resume_templates (name, description, industry_tags, preview_url, template_file, is_active)
                    VALUES (:name, :description, :industry_tags::jsonb, :preview_url, :template_file, :is_active)
                """),
                {
                    "name": t["name"],
                    "description": t.get("description", ""),
                    "industry_tags": json.dumps(t.get("industry_tags", [])),
                    "preview_url": t.get("preview_url", ""),
                    "template_file": tfile,
                    "is_active": t.get("is_active", True),
                },
            )
            inserted += 1
        except Exception as ex:
            logger.debug("resume_templates seed skip %s: %s", tfile, ex)
    if inserted:
        logger.info("Seeded %d new resume templates (total now: %d)", inserted, count + inserted)


@app.on_event("startup")
def _startup_check():
    require_production_secret()

    # Ensure interactive assessment tables exist
    try:
        from backend.app.routers.interactive_assess import ensure_assessment_tables
        from backend.app.db.session import SessionLocal
        db = SessionLocal()
        try:
            ensure_assessment_tables(db)
            db.commit()
        finally:
            db.close()
        logger.info("Assessment tables ensured.")
    except Exception as exc:
        logger.warning("Could not ensure assessment tables: %s", exc)

    # Seed roles + skills from JSON files if DB tables are empty
    try:
        from backend.app.db.session import SessionLocal as _SL
        _db = _SL()
        try:
            _seed_roles_and_skills(_db)
            _db.commit()
        finally:
            _db.close()
    except Exception as exc:
        logger.warning("Could not seed roles/skills: %s", exc)

    # Seed resume templates from JSON if DB has fewer than expected
    try:
        from backend.app.db.session import SessionLocal as _SL2
        _db2 = _SL2()
        try:
            _seed_resume_templates(_db2)
            _db2.commit()
        finally:
            _db2.close()
    except Exception as exc:
        logger.warning("Could not seed resume templates: %s", exc)

    global _SCHEMA_HEALTH
    _SCHEMA_HEALTH = _run_schema_health_check()
    if not _SCHEMA_HEALTH.get("ok", False):
        logger.warning("Schema health check warnings: %s", _SCHEMA_HEALTH.get("missing"))


import os as _os
import re as _re

# Always allow known SkillSight frontends so production (e.g. Railway) env override doesn't block Pages.
_CORS_KNOWN_FRONTENDS = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "https://skillsight-230.pages.dev",
    "https://skillsight.pages.dev",
]
_env_origins = [
    o.strip()
    for o in _os.getenv("CORS_ALLOWED_ORIGINS", "").strip().split(",")
    if o.strip()
]
_CORS_ORIGINS_RAW = list(dict.fromkeys(_env_origins + _CORS_KNOWN_FRONTENDS)) if _env_origins else _CORS_KNOWN_FRONTENDS
_CORS_PATTERNS = [
    _re.compile(r"^https://[a-z0-9-]+\.skillsight-\d+\.pages\.dev$"),
    _re.compile(r"^https://skillsight-\d+\.pages\.dev$"),
    _re.compile(r"^https://skillsight\.pages\.dev$"),
]
if not _os.getenv("SKILLSIGHT_ENV", "").strip().lower() in ("production", "prod"):
    _CORS_PATTERNS.append(_re.compile(r"^https://[a-z0-9-]+\.trycloudflare\.com$"))


def _origin_allowed(origin: str) -> bool:
    if not origin:
        return False
    if origin in _CORS_ORIGINS_RAW:
        return True
    return any(p.match(origin) for p in _CORS_PATTERNS)


_CORS_ALLOWED_HEADERS = [
    "Content-Type", "Authorization", "X-Requested-With", "X-Purpose",
    "X-Idempotency-Key", "X-Model-Version", "X-Rubric-Version", "Accept", "Origin",
]

# Middleware ordering: added in reverse (last = outermost in Starlette).
# CORSMiddleware MUST be outermost so CORS headers are always present,
# even when inner middlewares raise unhandled errors.
app.add_middleware(AuditMiddleware)
if _parse_bool_env("RATE_LIMIT_ENABLED"):
    app.add_middleware(RateLimitMiddleware)
app.add_middleware(
    CORSMiddleware,
    allow_origin_regex=r"^https://(skillsight-\d+\.pages\.dev|[a-z0-9-]+\.skillsight-\d+\.pages\.dev|skillsight\.pages\.dev)$",
    allow_origins=_CORS_ORIGINS_RAW,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=_CORS_ALLOWED_HEADERS,
    max_age=3600,
)

_main_logger = logging.getLogger("skillsight.main")

@app.exception_handler(Exception)
async def _global_exception_handler(request: Request, exc: Exception):
    _main_logger.exception("Unhandled exception on %s %s", request.method, request.url.path)
    return JSONResponse(
        status_code=500,
        content={"detail": "Internal server error"},
    )

@app.get("/")
def root():
    return {
        "service": "SkillSight API",
        "version": "0.1.3-fix-template-apply",
        "status": "ok",
        "docs": "/docs",
        "health": "/health",
    }


@app.get("/health/templates")
def health_templates():
    """Diagnostic: check template files, python-docx, and full apply flow."""
    import io
    from pathlib import Path
    base = Path(__file__).resolve().parents[1] / "data" / "templates"
    result: dict = {"templates_dir": str(base), "exists": base.exists()}
    if base.exists():
        result["files"] = sorted(f.name for f in base.iterdir() if f.is_file())
    else:
        result["files"] = []
    try:
        from docx import Document
        result["python_docx"] = "ok"
        if result["files"]:
            test_path = base / result["files"][0]
            doc = Document(str(test_path))
            result["test_open"] = f"ok ({len(doc.paragraphs)} paragraphs)"
            for p in doc.paragraphs:
                if "{{ RESUME_CONTENT }}" in p.text or "{{ CONTENT }}" in p.text:
                    p.runs[0].text = p.text.replace("{{ RESUME_CONTENT }}", "Test content here").replace("{{ CONTENT }}", "Test content here")
                    for run in p.runs[1:]:
                        run.text = ""
            buf = io.BytesIO()
            doc.save(buf)
            result["test_save"] = f"ok ({buf.tell()} bytes)"
    except Exception as e:
        import traceback
        result["python_docx"] = f"error: {type(e).__name__}: {e}"
        result["traceback"] = traceback.format_exc()

    try:
        from sqlalchemy import text as sa_text
        from backend.app.db.session import engine as _engine
        with _engine.connect() as conn:
            r = conn.execute(sa_text("SELECT COUNT(*) FROM resume_templates")).scalar()
            result["resume_templates_table"] = f"ok ({r} rows)"
    except Exception as e:
        result["resume_templates_table"] = f"error: {type(e).__name__}: {e}"

    try:
        from backend.app.services.resume_template_service import apply_template as _ta
        from backend.app.db.session import engine as _engine2
        from sqlalchemy.orm import Session as _Sess
        with _Sess(_engine2) as db:
            doc_bytes = _ta(db, review_id="test", template_id="__professional_classic", resume_content="John Doe\nSoftware Engineer\nExperience: 5 years in Python development.")
            result["test_apply"] = f"ok ({len(doc_bytes)} bytes)"
    except Exception as e:
        import traceback
        result["test_apply"] = f"error: {type(e).__name__}: {e}"
        result["test_apply_traceback"] = traceback.format_exc()

    return result


@app.get("/__routes")
def __routes():
    if _is_production():
        raise HTTPException(status_code=404, detail="Not found")
    return JSONResponse([
        {"path": r.path, "name": r.name, "methods": sorted(list(getattr(r, "methods", []) or []))}
        for r in app.router.routes
    ])

@app.get("/health")
def health():
    # Re-check schema on each request so DB fixes (migrations / manual SQL) show up without redeploy.
    live = _run_schema_health_check()
    return {"status": "ok", "ok": True, "schema": live}


@app.get("/health/schema")
def health_schema():
    return _SCHEMA_HEALTH


def _pg_health() -> dict:
    """Check PostgreSQL connectivity."""
    try:
        with SessionLocal() as db:
            db.execute(text("SELECT 1"))
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _redis_health() -> dict:
    """Check Redis connectivity."""
    try:
        import os as _os
        import redis
        host = _os.getenv("REDIS_HOST", "localhost")
        port = int(_os.getenv("REDIS_PORT", "6379"))
        password = _os.getenv("REDIS_PASSWORD") or _os.getenv("REDIS_PASSWORD_ENV")
        r = redis.Redis(host=host, port=port, password=password, db=0, socket_connect_timeout=2)
        r.ping()
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/readyz")
def readyz():
    """Readiness: health + PostgreSQL, Redis, Qdrant connectivity."""
    from backend.app.vector_store import qdrant_health
    qdrant = qdrant_health()
    pg = _pg_health()
    redis_status = _redis_health()
    ok = qdrant.get("ok", False) and pg.get("ok", False) and redis_status.get("ok", False)
    return {
        "status": "ok" if ok else "degraded",
        "ok": ok,
        "postgres": pg,
        "redis": redis_status,
        "qdrant": qdrant,
    }


@app.get("/debug/routes_count")
def debug_routes_count():
    """Debug endpoint to count routes."""
    if _is_production():
        raise HTTPException(status_code=404, detail="Not found")
    return {
        "total_routes": len(app.routes),
        "routes_preview": [r.path for r in app.routes if hasattr(r, 'path')][:20],
        "version": "v2_with_interactive"
    }


@app.get("/api/overview")
def api_overview():
    """
    API Overview - List all available endpoints grouped by feature.
    """
    return {
        "version": "0.2.0",
        "mvp_features": {
            "document_management": {
                "POST /documents/upload": "Upload TXT/DOCX/PDF/Images/Video with consent",
                "POST /documents/upload_multimodal": "Upload any supported file with multimodal parsing",
                "GET /documents": "List uploaded documents",
                "GET /documents/{doc_id}": "Get document details",
                "GET /documents/{doc_id}/chunks": "Get document chunks",
                "POST /documents/{doc_id}/reindex": "Trigger re-indexing",
            },
            "evidence_search": {
                "POST /search/evidence_vector": "Vector search for skill evidence (Decision 1)",
                "POST /search/evidence_keyword": "Keyword fallback search",
            },
            "ai_assessment": {
                "POST /ai/demonstration": "Skill demonstration classification (Decision 2)",
                "POST /ai/proficiency": "Proficiency level assessment (Decision 3)",
            },
            "role_readiness": {
                "POST /assess/role_readiness": "Role readiness assessment (Decision 4)",
            },
            "action_recommendations": {
                "POST /actions/recommend": "Action cards for skill gaps (Decision 5)",
                "GET /actions/templates": "List action templates",
            },
            "consent_management": {
                "POST /consent/grant": "Grant consent for document",
                "POST /consent/revoke": "Revoke consent (cascade delete)",
                "GET /consent/status/{doc_id}": "Check consent status",
            },
            "background_jobs": {
                "GET /jobs": "List background jobs",
                "GET /jobs/{job_id}": "Get job status",
                "POST /jobs/{job_id}/retry": "Retry failed job",
                "POST /jobs/enqueue/{doc_id}": "Enqueue new job",
                "GET /jobs/queue/status": "Redis queue status",
            },
            "direct_embedding": {
                "POST /chunks/embed/{doc_id}": "Sync embedding (no Redis)",
            },
        },
        "interactive_assessments": {
            "communication_kira_style": {
                "POST /interactive/communication/start": "Start video response session (random topic)",
                "POST /interactive/communication/submit": "Submit speech transcript for evaluation",
                "description": "Kira-style assessment: random topic, 30/60/90 sec response, retry support",
            },
            "programming_leetcode_style": {
                "POST /interactive/programming/start": "Start coding challenge (easy/medium/hard)",
                "POST /interactive/programming/submit": "Submit code solution for evaluation",
                "description": "LeetCode-style challenges with auto-generated problems",
            },
            "writing_timed": {
                "POST /interactive/writing/start": "Start timed writing session",
                "POST /interactive/writing/submit": "Submit essay for evaluation",
                "description": "Timed writing with anti-copy protection (300-500 words, 30 min)",
            },
            "session_management": {
                "GET /interactive/sessions/{session_id}": "Get session details and attempts",
                "GET /interactive/sessions/user/{user_id}": "Get all sessions for a user",
                "GET /interactive/sessions/{session_id}/consistency": "Calculate consistency across retries",
            },
        },
        "multimodal_support": {
            "supported_formats": {
                "documents": [".txt", ".docx", ".pdf", ".pptx"],
                "images": [".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".gif"],
                "video_audio": [".mp4", ".webm", ".mov", ".avi", ".mp3", ".wav", ".m4a"],
                "code": [".py", ".js", ".ts", ".java", ".cpp", ".c", ".go", ".rs"],
            },
            "processing": {
                "images": "OCR text extraction (pytesseract/easyocr) + vision model support",
                "video_audio": "Whisper transcription + metadata extraction",
                "presentations": "Slide text + speaker notes extraction",
            },
        },
        "data_management": {
            "skills": "GET/POST /skills, /skills/search, /skills/import",
            "roles": "GET/POST /roles, /roles/{id}, /roles/{id}/requirements",
            "courses": "GET /courses",
        },
    }

@app.get("/stats")
def stats(db: Session = Depends(get_db)):
    """
    Lightweight counts for demo smoke-check.
    Defensive: should not fail if table names change.
    """
    tables = db.execute(
        text("""
            SELECT table_name
            FROM information_schema.tables
            WHERE table_schema = 'public'
              AND table_type = 'BASE TABLE'
            ORDER BY table_name
        """)
    ).fetchall()
    return {
        "status": "ok",
        "public_tables": [r[0] for r in tables],
        "public_table_count": len(tables),
    }

@app.get("/schemas/summary")
def schemas_summary():
    """
    Returns the JSONSchema filenames generated under packages/schemas/.
    This is a simple discoverability endpoint for the MVP.
    """
    return {
        "Skill": "packages/schemas/Skill.schema.json",
        "Role": "packages/schemas/Role.schema.json",
        "EvidencePointer": "packages/schemas/EvidencePointer.schema.json",
        "AuditLog": "packages/schemas/AuditLog.schema.json",
        "ConsentRecord": "packages/schemas/ConsentRecord.schema.json",
    }

from backend.app.routers.skills import router as skills_router
from backend.app.routers.roles import router as roles_router
from backend.app.routers.documents import router as documents_router
from backend.app.routers.chunks import router as chunks_router
from backend.app.routers.consents import router as consents_router
from backend.app.routers.jobs import router as jobs_router
from backend.app.routers.courses import router as courses_router
from backend.app.routers.assessments import router as assessments_router
from backend.app.routers.proficiency import router as proficiency_router
from backend.app.routers.ai import router as ai_router
from backend.app.routers.assess import router as assess_router
from backend.app.routers.search import router as search_router
from backend.app.routers.actions import router as actions_router
from backend.app.routers.interactive_assess import router as interactive_router
from backend.app.routers.auth import router as auth_router
# BFF tier routers (P2)
from backend.app.routers.bff_student import router as bff_student_router
from backend.app.routers.resume_review import router as resume_review_router
from backend.app.routers.bff_staff import router as bff_staff_router
from backend.app.routers.bff_programme import router as bff_programme_router
from backend.app.routers.bff_admin import router as bff_admin_router


ROUTERS = [skills_router, roles_router]

for r in ROUTERS:
    app.include_router(r)
app.include_router(documents_router)
app.include_router(chunks_router)
app.include_router(consents_router)
app.include_router(jobs_router)
app.include_router(courses_router)
app.include_router(assessments_router)
app.include_router(proficiency_router)
app.include_router(ai_router)
app.include_router(assess_router)
app.include_router(search_router)
app.include_router(actions_router)
app.include_router(interactive_router)
app.include_router(auth_router)
# BFF tier (P2)
app.include_router(bff_student_router)
app.include_router(resume_review_router, prefix="/bff/student")
app.include_router(bff_staff_router)
app.include_router(bff_programme_router)
app.include_router(bff_admin_router)
# NOTE: /skills endpoint is now handled by skills_router from routers/skills.py
# Removed duplicate endpoint to avoid route conflicts
