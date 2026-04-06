"""
Document Parsers for SkillSight
- TXT, DOCX, PDF parsing with chunking
"""
import hashlib
import io
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


def _make_snippet(text: str, max_len: int = 220) -> str:
    """Create a snippet from text."""
    t = (text or "").strip().replace("\n", " ")
    t = re.sub(r"\s+", " ", t)
    if len(t) <= max_len:
        return t
    return t[:max_len] + "..."


def _compute_hash(text: str) -> str:
    """Compute SHA-256 hash of text."""
    return hashlib.sha256((text or "").encode("utf-8")).hexdigest()


def _coalesce_short_paragraphs(
    parts: List[str],
    min_chunk_len: int = 50,
    max_chunk_len: int = 2000,
) -> List[str]:
    """
    Merge consecutive short segments so fewer sentences are dropped before indexing.
    Aligns with protocol min_chunk_length while preserving content (resume-friendly).
    """
    items = [p.strip() for p in parts if p and p.strip()]
    if not items:
        return []
    out: List[str] = []
    buf = items[0]
    for x in items[1:]:
        if len(buf) >= min_chunk_len:
            if len(buf) + 1 + len(x) <= max_chunk_len:
                buf = f"{buf} {x}"
            else:
                out.append(buf)
                buf = x
        else:
            if len(buf) + 1 + len(x) <= max_chunk_len:
                buf = f"{buf} {x}"
            else:
                out.append(buf)
                buf = x
    out.append(buf)
    # Merge trailing fragment into previous if still below min (single small tail)
    while len(out) >= 2 and len(out[-1]) < min_chunk_len:
        if len(out[-2]) + 1 + len(out[-1]) <= max_chunk_len:
            out[-2] = f"{out[-2]} {out[-1]}"
            out.pop()
        else:
            break
    return out


# ====================
# TXT Parser
# ====================
def parse_txt_to_chunks(
    content: str,
    min_chunk_len: int = 50,
) -> List[Dict[str, Any]]:
    """
    Parse plain text into chunks by double newlines.
    Returns list of chunk dicts with char_start, char_end, chunk_text, snippet, quote_hash.
    """
    if not content:
        return []
    
    text = content.replace("\r\n", "\n").replace("\r", "\n")
    raw_parts = text.split("\n\n")
    stripped_parts = [p.strip() for p in raw_parts if p.strip()]
    merged_texts = _coalesce_short_paragraphs(stripped_parts, min_chunk_len)
    chunks = []
    cursor = 0
    idx = 0
    for part_strip in merged_texts:
        pos = text.find(part_strip, cursor)
        if pos == -1:
            pos = cursor
        char_start = pos
        char_end = pos + len(part_strip)
        chunks.append({
            "idx": idx,
            "char_start": char_start,
            "char_end": char_end,
            "chunk_text": part_strip,
            "snippet": _make_snippet(part_strip),
            "quote_hash": _compute_hash(part_strip),
            "section_path": None,
            "page_start": None,
            "page_end": None,
        })
        cursor = char_end
        idx += 1
    return chunks


# ====================
# DOCX Parser
# ====================
def _parse_docx_document_to_chunks(doc: Any, min_chunk_len: int = 50) -> List[Dict[str, Any]]:
    """Shared DOCX chunking: merge consecutive short body paragraphs (non-heading)."""
    chunks: List[Dict[str, Any]] = []
    idx = 0
    char_cursor = 0
    current_section: List[str] = []
    pending: List[str] = []

    def flush_pending() -> None:
        nonlocal pending, idx, char_cursor, chunks, current_section
        if not pending:
            return
        merged = " ".join(pending)
        pending = []
        section_path = " > ".join(current_section) if current_section else None
        chunks.append({
            "idx": idx,
            "char_start": char_cursor,
            "char_end": char_cursor + len(merged),
            "chunk_text": merged,
            "snippet": _make_snippet(merged),
            "quote_hash": _compute_hash(merged),
            "section_path": section_path,
            "page_start": None,
            "page_end": None,
        })
        char_cursor += len(merged) + 1
        idx += 1

    def append_chunk(chunk_text: str) -> None:
        nonlocal idx, char_cursor, chunks, current_section
        section_path = " > ".join(current_section) if current_section else None
        chunks.append({
            "idx": idx,
            "char_start": char_cursor,
            "char_end": char_cursor + len(chunk_text),
            "chunk_text": chunk_text,
            "snippet": _make_snippet(chunk_text),
            "quote_hash": _compute_hash(chunk_text),
            "section_path": section_path,
            "page_start": None,
            "page_end": None,
        })
        char_cursor += len(chunk_text) + 1
        idx += 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        is_heading = style_name.lower().startswith("heading")

        if is_heading:
            flush_pending()
            level_match = re.search(r"(\d+)", style_name)
            level = int(level_match.group(1)) if level_match else 1
            while len(current_section) >= level:
                current_section.pop()
            current_section.append(text)
            append_chunk(text)
            continue

        if len(text) >= min_chunk_len:
            flush_pending()
            append_chunk(text)
            continue

        pending.append(text)
        joined = " ".join(pending)
        if len(joined) >= min_chunk_len:
            flush_pending()

    flush_pending()
    return chunks


def parse_docx_to_chunks(
    file_path: str,
    min_chunk_len: int = 50,
) -> List[Dict[str, Any]]:
    """
    Parse DOCX file into chunks by paragraphs.
    Preserves heading hierarchy in section_path.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

    doc = Document(file_path)
    return _parse_docx_document_to_chunks(doc, min_chunk_len)


def parse_docx_bytes_to_chunks(
    file_bytes: bytes,
    min_chunk_len: int = 50,
) -> List[Dict[str, Any]]:
    """
    Parse DOCX from bytes into chunks.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

    doc = Document(io.BytesIO(file_bytes))
    return _parse_docx_document_to_chunks(doc, min_chunk_len)


# ====================
# PDF Parser
# ====================
def parse_pdf_to_chunks(
    file_path: str,
    min_chunk_len: int = 50,
) -> List[Dict[str, Any]]:
    """
    Parse PDF file into chunks by pages/paragraphs.
    Preserves page numbers.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF parsing. Install with: pip install pymupdf")
    
    doc = fitz.open(file_path)
    chunks = []
    idx = 0
    char_cursor = 0
    
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text")
        if not text.strip():
            continue
        
        paragraphs = text.split("\n\n")
        normalized = [re.sub(r"\s+", " ", p.strip()) for p in paragraphs if p.strip()]
        merged_paras = _coalesce_short_paragraphs(normalized, min_chunk_len)

        for para_text in merged_paras:
            chunks.append({
                "idx": idx,
                "char_start": char_cursor,
                "char_end": char_cursor + len(para_text),
                "chunk_text": para_text,
                "snippet": _make_snippet(para_text),
                "quote_hash": _compute_hash(para_text),
                "section_path": None,
                "page_start": page_num,
                "page_end": page_num,
            })
            char_cursor += len(para_text) + 2
            idx += 1
    
    doc.close()
    return chunks


def parse_pdf_bytes_to_chunks(
    file_bytes: bytes,
    min_chunk_len: int = 50,
) -> List[Dict[str, Any]]:
    """
    Parse PDF from bytes into chunks.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF parsing. Install with: pip install pymupdf")
    
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    chunks = []
    idx = 0
    char_cursor = 0
    
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text")
        if not text.strip():
            continue
        
        paragraphs = text.split("\n\n")
        normalized = [re.sub(r"\s+", " ", p.strip()) for p in paragraphs if p.strip()]
        merged_paras = _coalesce_short_paragraphs(normalized, min_chunk_len)

        for para_text in merged_paras:
            chunks.append({
                "idx": idx,
                "char_start": char_cursor,
                "char_end": char_cursor + len(para_text),
                "chunk_text": para_text,
                "snippet": _make_snippet(para_text),
                "quote_hash": _compute_hash(para_text),
                "section_path": None,
                "page_start": page_num,
                "page_end": page_num,
            })
            char_cursor += len(para_text) + 2
            idx += 1
    
    doc.close()
    return chunks


# ====================
# Unified Parser
# ====================
def parse_file_to_chunks(
    file_path: str = None,
    file_bytes: bytes = None,
    filename: str = None,
    min_chunk_len: int = 50,
) -> List[Dict[str, Any]]:
    """
    Unified parser that detects file type and parses accordingly.
    
    Args:
        file_path: Path to file (for file-based parsing)
        file_bytes: File content as bytes (for in-memory parsing)
        filename: Filename to detect extension (required if using file_bytes)
        min_chunk_len: Minimum chunk length
    
    Returns:
        List of chunk dictionaries
    """
    if file_path:
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".txt" or ext == ".md" or ext == ".markdown":
            content = Path(file_path).read_text(encoding="utf-8", errors="ignore")
            return parse_txt_to_chunks(content, min_chunk_len)
        elif ext == ".csv":
            content = Path(file_path).read_text(encoding="utf-8", errors="ignore")
            return parse_txt_to_chunks(content, min_chunk_len)
        elif ext == ".docx":
            return parse_docx_to_chunks(file_path, min_chunk_len)
        elif ext == ".pdf":
            return parse_pdf_to_chunks(file_path, min_chunk_len)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    elif file_bytes and filename:
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == ".txt" or ext == ".md" or ext == ".markdown":
            try:
                content = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                content = file_bytes.decode("utf-8", errors="ignore")
            return parse_txt_to_chunks(content, min_chunk_len)
        elif ext == ".csv":
            try:
                content = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                content = file_bytes.decode("utf-8", errors="ignore")
            return parse_txt_to_chunks(content, min_chunk_len)
        elif ext == ".docx":
            return parse_docx_bytes_to_chunks(file_bytes, min_chunk_len)
        elif ext == ".pdf":
            return parse_pdf_bytes_to_chunks(file_bytes, min_chunk_len)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    else:
        raise ValueError("Either file_path or (file_bytes + filename) must be provided")
