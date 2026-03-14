#!/usr/bin/env python3
"""
Seed resume_templates table and create minimal DOCX template files.
Run from repo root: PYTHONPATH=backend python scripts/seed_resume_templates.py
Requires: DATABASE_URL, and table resume_templates exists (run alembic upgrade head first).
"""
from __future__ import annotations

import json
import sys
import uuid
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
BACKEND = REPO_ROOT / "backend"
if str(BACKEND) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

def main():
    from sqlalchemy import text
    from backend.app.db.session import engine

    seeds_file = REPO_ROOT / "backend" / "data" / "seeds" / "resume_templates.json"
    templates_dir = REPO_ROOT / "backend" / "data" / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)

    if not seeds_file.exists():
        print(f"Seeds file not found: {seeds_file}", file=sys.stderr)
        sys.exit(1)
    items = json.loads(seeds_file.read_text())
    if not items:
        print("No templates in JSON")
        return

    # Create minimal DOCX files with placeholder (python-docx)
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("python-docx not installed; skipping DOCX file creation.", file=sys.stderr)
        Document = None

    for item in items:
        tfile = (item.get("template_file") or "").strip()
        if not tfile:
            continue
        if tfile.endswith(".docx"):
            path = templates_dir / tfile
        else:
            path = templates_dir / f"{tfile}.docx"
        if Document and not path.exists():
            doc = Document()
            p = doc.add_paragraph()
            p.add_run("{{ RESUME_CONTENT }}")
            doc.save(str(path))
            print(f"  Created template file: {path.name}")

    with engine.connect() as conn:
        for item in items:
            name = item.get("name") or "Unnamed"
            existing = conn.execute(
                text("SELECT 1 FROM resume_templates WHERE name = :name LIMIT 1"),
                {"name": name},
            ).scalar()
            if existing:
                print(f"  Skip (exists): {name}")
                continue
            template_id = str(uuid.uuid4())
            description = (item.get("description") or "")[:2000]
            industry_tags = item.get("industry_tags")
            if isinstance(industry_tags, list):
                industry_tags = json.dumps(industry_tags)
            else:
                industry_tags = "[]"
            preview_url = (item.get("preview_url") or "")[:500]
            template_file = (item.get("template_file") or "")[:255]
            is_active = bool(item.get("is_active", True))
            conn.execute(
                text("""
                    INSERT INTO resume_templates (template_id, name, description, industry_tags, preview_url, template_file, is_active, created_at)
                    VALUES (:tid, :name, :description, :industry_tags::jsonb, :preview_url, :template_file, :is_active, NOW())
                """),
                {
                    "tid": template_id,
                    "name": name,
                    "description": description,
                    "industry_tags": industry_tags,
                    "preview_url": preview_url,
                    "template_file": template_file,
                    "is_active": is_active,
                },
            )
            print(f"  Inserted template: {name} ({template_id[:8]}...)")
        conn.commit()
    print("Resume templates seed done.")


if __name__ == "__main__":
    main()
