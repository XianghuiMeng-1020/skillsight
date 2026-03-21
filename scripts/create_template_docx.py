#!/usr/bin/env python3
"""
Create 4 professional resume template DOCX files under backend/data/templates/.
Each template has a distinct visual style with placeholder {{ RESUME_CONTENT }}.
"""
from pathlib import Path
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
except ImportError:
    print("python-docx is required: pip install python-docx", file=sys.stderr)
    sys.exit(1)

TEMPLATES_DIR = Path(__file__).resolve().parents[1] / "backend" / "data" / "templates"
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


def _set_cell_shading(cell, hex_color: str):
    from docx.oxml.ns import qn
    from lxml import etree
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shading = etree.SubElement(tc_pr, qn("w:shd"))
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")


def create_professional_classic():
    """Clean, traditional layout — finance / consulting / corporate."""
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    h = doc.add_heading("Professional Classic Resume", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.size = Pt(8)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("{{ RESUME_CONTENT }}")

    doc.save(str(TEMPLATES_DIR / "professional_classic.docx"))
    print("  Created: professional_classic.docx")


def create_modern_tech():
    """Contemporary design — technology / engineering / software."""
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Segoe UI"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    _set_cell_shading(left_cell, "1e293b")
    p = left_cell.paragraphs[0]
    run = p.add_run("MODERN TECH")
    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    run.font.size = Pt(18)
    run.font.bold = True
    p2 = left_cell.add_paragraph()
    run2 = p2.add_run("Resume Template")
    run2.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    run2.font.size = Pt(11)

    _set_cell_shading(right_cell, "f1f5f9")
    pr = right_cell.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run3 = pr.add_run("SkillSight")
    run3.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    run3.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("{{ RESUME_CONTENT }}")

    doc.save(str(TEMPLATES_DIR / "modern_tech.docx"))
    print("  Created: modern_tech.docx")


def create_creative_portfolio():
    """Stylish layout — marketing / design / creative roles."""
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x37, 0x37, 0x37)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    h = doc.add_heading("", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run("Creative")
    run.font.color.rgb = RGBColor(0x7c, 0x3a, 0xed)
    run.font.size = Pt(28)
    run2 = h.add_run(" Portfolio")
    run2.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)
    run2.font.size = Pt(28)

    p = doc.add_paragraph()
    run = p.add_run("▬" * 8)
    run.font.color.rgb = RGBColor(0x7c, 0x3a, 0xed)
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("{{ RESUME_CONTENT }}")

    doc.save(str(TEMPLATES_DIR / "creative_portfolio.docx"))
    print("  Created: creative_portfolio.docx")


def create_academic_research():
    """Structured format — research / academia / education."""
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    h = doc.add_heading("Curriculum Vitae", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0a, 0x2a, 0x4a)
        run.font.size = Pt(24)
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 50)
    run.font.color.rgb = RGBColor(0x0a, 0x2a, 0x4a)
    run.font.size = Pt(8)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("{{ RESUME_CONTENT }}")

    doc.save(str(TEMPLATES_DIR / "academic_research.docx"))
    print("  Created: academic_research.docx")


if __name__ == "__main__":
    create_professional_classic()
    create_modern_tech()
    create_creative_portfolio()
    create_academic_research()
    print("All template DOCX files created.")
