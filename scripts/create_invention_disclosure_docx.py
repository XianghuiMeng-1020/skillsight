#!/usr/bin/env python3
"""
Generate the SkillSight Invention Disclosure form as a Word document.
Run from project root: python scripts/create_invention_disclosure_docx.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 11)
    p.space_after = Pt(6)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    p.paragraph_format.space_after = Pt(3)
    return p

def add_fill_line(doc, label, content):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run("\n" + content)
    p.paragraph_format.space_after = Pt(6)
    return p

def main():
    doc = Document()
    doc.add_heading("SkillSight – Invention Disclosure Form", 0)

    # 2. TITLE OF INVENTION
    add_heading(doc, "2. TITLE OF INVENTION", 1)
    add_para(doc, "Note: The title should describe what the invention does, but not how it is made or how it works.")
    add_para(doc, "Evidence-Based Skill Assessment and Job-Readiness Transparency System for Higher Education")

    # 3. BRIEF DESCRIPTION
    add_heading(doc, "3. BRIEF DESCRIPTION / ABSTRACT OF THE INVENTION", 1)

    add_para(doc, "a. Provide a short general layperson's overview of the invention and how it works.")
    add_para(doc, "SkillSight is a software system that allows university students to upload their academic work "
        "(documents, code, presentations, videos, etc.) and automatically evaluates what professional skills they "
        "have demonstrated, how proficient they are, and how ready they are for specific job roles. The system uses "
        "artificial intelligence to read and understand student work, find evidence of skills, assess proficiency "
        "levels against defined rubrics, and compare the student's skill profile against real job market requirements. "
        "It then generates personalized action plans to help students close any skill gaps. Every conclusion is "
        "traceable back to specific evidence in the student's submitted work, ensuring full transparency and explainability.")

    add_para(doc, "b. If the invention solves an existing problem or fulfil a long-felt need, what problem does it solve?")
    add_para(doc, "Currently, there is no systematic, evidence-based method for university students to demonstrate "
        "their skills to potential employers, nor for universities to objectively measure how well their curricula "
        "prepare students for the job market. Traditional transcripts only show courses taken and grades earned, but "
        "do not reflect the specific skills a student has developed or their proficiency levels. Employers struggle "
        "to verify candidate skills beyond self-reported claims on resumes. Universities lack data-driven tools to "
        "assess the alignment between their academic programmes and evolving industry demands. Additionally, existing "
        "AI-based assessment tools often produce \"black box\" results without traceable evidence, raising concerns "
        "about reliability and accountability.")

    add_para(doc, "c. How does the invention solve the problem?")
    add_para(doc, "The invention implements a five-stage AI decision pipeline that processes student-submitted artifacts: "
        "(1) Evidence Retrieval: Uses vector embeddings and semantic search to find relevant evidence chunks from "
        "uploaded documents that relate to specific skills. (2) Demonstration Assessment: An AI model determines "
        "whether a skill is \"demonstrated,\" \"mentioned,\" or has \"not enough information,\" with mandatory citation "
        "of specific evidence chunks. (3) Proficiency Assessment: Evaluates the proficiency level (0-3: Novice to "
        "Advanced) against structured rubrics, with each level requiring specific criteria to be met. (4) Role "
        "Readiness: Compares the student's assessed skill profile against job role requirements sourced from real job "
        "market data, producing readiness scores. (5) Action Recommendations: Generates personalized, actionable "
        "guidance for closing identified skill gaps, including specific artifacts to produce and how to verify "
        "completion. A critical \"fail-closed\" design ensures the system refuses to make claims when evidence is "
        "insufficient, preventing false positives and hallucinated assessments. All decisions are fully auditable "
        "with traceable evidence pointers back to the original source material.")

    add_para(doc, "d. Is it related to a new product, process, or composition of matter? Or is it a new use for or an "
        "improvement to an existing product, process or composition of matter? If so, what are the advantages of the "
        "invention over the existing ones, if any?")
    add_para(doc, "This is a new product. While individual components such as AI-based text analysis, vector search, "
        "and skill assessment exist independently, SkillSight uniquely integrates them into a unified, evidence-traceable "
        "pipeline specifically designed for higher education skill transparency. Advantages over existing solutions: "
        "Unlike traditional e-portfolio systems (e.g., Mahara, Portfolium), SkillSight provides automated, AI-driven "
        "skill extraction and assessment rather than relying on manual self-reporting. Unlike generic AI assessment "
        "tools, SkillSight enforces a strict \"no pointer, no claim\" policy with a fail-closed refusal mechanism, "
        "preventing AI hallucination. Unlike employer-facing skill verification platforms (e.g., LinkedIn Skills "
        "Assessment), SkillSight evaluates skills based on authentic student work products rather than standardized "
        "quizzes. The system provides a complete, auditable decision chain from raw evidence to job readiness, which no "
        "existing product offers in the higher education context. Multi-modal support (documents, code, video, images) "
        "with privacy-by-design consent management ensures comprehensive and compliant skill assessment.")

    add_para(doc, "e. Identify the elements that are considered novel and unobvious.")
    novel_items = [
        "1. Fail-Closed Refusal Mechanism with Structured Refusal Responses: The pipeline enforces mandatory refusal "
        "when evidence is insufficient, returning structured refusal objects (code, message, next_step). The system "
        "never makes a positive skill claim without verifiable evidence pointers.",
        "2. Five-Stage Cascading Decision Pipeline with Evidence Traceability: Every conclusion includes specific "
        "chunk IDs, character offsets, and document sections that support it, with end-to-end traceability.",
        "3. Reliability-Tiered Evidence Assessment: Retrieval pipeline computes reliability (high/medium/low) from "
        "score distributions, gap analysis, and reranker stability.",
        "4. Cross-Document Skill Level Aggregation with Conflict Detection: Consistency checks, mutual conflict "
        "detection, minimum evidence thresholds; fail-closed to level 0 and flag for human review on conflict.",
        "5. Consent-Gated Processing Pipeline with Cascade Deletion: Evidence retrieval and assessment gated by "
        "consent; revoking consent triggers full cascade deletion across all derived data.",
        "6. Course-Skill-Role Triangulation: Three-way mapping between courses, skills, and job roles for transparent "
        "visibility into curriculum-to-career alignment.",
    ]
    for item in novel_items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()

    # 4. TECHNICAL DESCRIPTION
    add_heading(doc, "4. TECHNICAL DESCRIPTION OF THE INVENTION", 1)
    add_para(doc, "Note: Please provide any relevant sketches, diagrams, drawings, photographs or other illustrative "
        "material. Description may reference a separate document such as a publication, meeting abstract, and "
        "manuscript in preparation, preprint or report.")
    tech_desc = """System Architecture:
SkillSight is a full-stack web application. Frontend: Next.js (React) with student dashboards, document upload, skill visualization, interactive assessments. Backend: FastAPI (Python) with 40+ endpoints and role-based BFF layers. Data: PostgreSQL (pgvector), Qdrant (vector search), Redis (task queue). AI/ML: OpenAI or Ollama for LLM; Sentence Transformers or OpenAI for embeddings; Whisper for transcription; OCR for images.

Core Decision Pipeline (Decisions 1–5):
Decision 1 – Evidence Retrieval: Documents parsed into chunks (800 chars, 100 overlap) with character offsets; embedded and stored in Qdrant. Query vector from skill definition; ANN search; pre/post threshold filtering; reliability from score distributions and reranker stability.
Decision 2 – Demonstration Assessment: Retrieved chunks passed to LLM with strict rules (no fabrication, mandatory chunk_id citation, refusal when insufficient). Post-LLM guardrails validate JSON, chunk_ids, and evidence lists. Output: label, evidence_chunk_ids, rationale, refusal_reason.
Decision 3 – Proficiency Assessment: Skill-specific rubrics (level 0–3); LLM cites criteria_ids. Output: level, matched_criteria, evidence_chunk_ids.
Decision 4 – Role Readiness: Compares assessed skills to job role requirements; each skill: meet / needs_strengthening / missing_proof.
Decision 5 – Action Recommendations: Action cards (what_to_do, artifact, how_verified, based_on) and links to learning resources.

Skill Level Aggregation: Cross-document majority voting, consistency ratio (80%), mutual conflict detection, min evidence >= 2 for high reliability; changes in skill_assessment_snapshots and change_log_events.

Interactive Assessments: Communication (Kira-style video), Programming (LeetCode-style), Writing (timed, anti-plagiarism).

Privacy: Consent-gated processing; cascade deletion on revoke; full audit logging.

Attachment if any: ☒ Data  ☒ Others (source code repository and API documentation)"""
    add_para(doc, tech_desc)

    # 5. GRANTS
    add_heading(doc, "5. GRANTS / FINANCIAL SUPPORT / CONTRACTUAL OBLIGATION", 1)
    add_para(doc, "Name of Sponsor: [To be completed by inventor]")
    add_para(doc, "Grant/contract no.: [To be completed by inventor]")
    add_para(doc, "Project Title and Project End Date: [To be completed by inventor]")
    add_para(doc, "Name of Project-In-Charge: [To be completed by inventor]")
    add_para(doc, "Funding for Patent Costs: ☐ No, no funding.  ☐ Yes, percentage or amount covered: ___")

    # 6. PUBLIC DISCLOSURE
    add_heading(doc, "6. PUBLIC DISCLOSURE", 1)
    add_para(doc, "☐ No, it has not been made available to the public.")
    add_para(doc, "☐ Yes, when and where was this invention disclosed: [To be completed by inventor]")
    add_para(doc, "If the invention has not been publicly disclosed, indicate when and where it will be disclosed: [To be completed by inventor]")

    # 7. STAGE OF DEVELOPMENT
    add_heading(doc, "7. STAGE OF DEVELOPMENT", 1)
    add_para(doc, "☐ Concept Only  ☐ Preliminary Data  ☐ Intermediate Data  ☒ Complete Proof of Concept  ☐ Prototype Available")
    add_para(doc, "Are there any limitations in the current invention or challenges that needed to be overcome before a practical product may be developed?")
    add_para(doc, "Yes. Key limitations: (1) Dependence on third-party LLM APIs (latency, cost, privacy). (2) Reranker is placeholder; production cross-encoder would improve precision. (3) Validated with HKU BASc(SDS) and BSc(IM); generalization to other institutions requires more mapping. (4) Scalability for concurrent users needs optimization. (5) Interactive assessment modules need psychometric validation.")
    add_para(doc, "Will you continue your work to overcome these limitations/challenges? What is your future development plan?")
    add_para(doc, "Yes. Plan includes: fine-tuning domain-specific models; production reranker; expanding course-skill-role mapping; validation studies vs expert assessment; institutional analytics dashboards; federated deployment options.")

    # 8. PATENT AND LITERATURE SEARCH
    add_heading(doc, "8. PATENT AND LITERATURE SEARCH", 1)
    add_para(doc, "a. Keyword(s) used in the patent or literature search")
    add_para(doc, "AI skill assessment education, evidence-based competency evaluation, automated portfolio assessment, "
        "skill gap analysis higher education, LLM educational assessment explainability, vector search skill matching, "
        "job readiness assessment system, AI hallucination prevention educational assessment, fail-closed AI evaluation, "
        "course-skill-job mapping platform")
    add_para(doc, "b. Prior arts (at least two):")
    add_para(doc, "1. Type: ☒ Market available product or technology. Portfolium (Instructure/Canvas) — e-portfolio with "
        "manual skill tagging. Difference: SkillSight performs automated AI-driven extraction and assessment from "
        "artifacts with fail-closed refusal and full evidence traceability; Portfolium does not. Not routine; contributes to advantageous effect.")
    add_para(doc, "2. Type: ☒ Market available product or technology. LinkedIn Skills Assessment — standardized quizzes. "
        "Difference: SkillSight assesses authentic work products with evidence traceability, proficiency grading, and "
        "job readiness mapping; LinkedIn uses quizzes without evidence chain. Not routine; contributes to advantageous effect.")

    # 9. POTENTIAL COMMERCIAL APPLICATIONS
    add_heading(doc, "9. POTENTIAL COMMERCIAL APPLICATIONS OF THE INVENTION", 1)
    add_para(doc, "a. Competing products: Portfolium, Handshake, Burning Glass/Lightcast, Coursera Skills Graph, Credly/Acclaim, PebblePad. None offer automated evidence-traceable AI skill assessment from authentic work with fail-closed pipeline and job readiness mapping.")
    add_para(doc, "b. Interested parties: University administrations (HKU, UGC institutions); EdTech (Instructure, Blackboard, D2L, Moodle); career platforms (Handshake); accreditation bodies (HKCAAVQ, QAA); EDB, UGC; enterprise learning (Workday, Degreed). Specific contacts to be provided by inventor.")
    add_para(doc, "c. Estimated cost: Cloud infra USD 500–2000/mo; LLM API USD 200–1000/mo; setup USD 10k–30k; maintenance USD 2k–5k/mo. Year-1 ~USD 50k–120k per institution. Self-hosted can reduce API costs 60–80%.")
    add_para(doc, "d. Most important markets: 1. Hong Kong  2. China  3. US")

    # 10. THIRD PARTY MATERIALS
    add_heading(doc, "10. THIRD PARTY MATERIALS", 1)
    add_para(doc, "☐ No third-party material or data was used.")
    add_para(doc, "☒ Third-party services and libraries: OpenAI API (LLM, embeddings, Whisper); Qdrant (Apache 2.0); "
        "Sentence Transformers (Apache 2.0); Ollama (MIT); FastAPI, SQLAlchemy, Next.js, React (MIT/BSD); LinkedIn job "
        "market data from public listings. No Material Transfer Agreements; components are commercial APIs or "
        "open-source with permissive licenses.")

    out_path = Path(__file__).resolve().parent.parent / "docs" / "SkillSight_Invention_Disclosure_Form.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    main()
