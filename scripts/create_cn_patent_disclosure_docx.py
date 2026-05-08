#!/usr/bin/env python3
"""
Generate a Chinese mainland patent disclosure document for SkillSight.

Run from project root:
    python scripts/create_cn_patent_disclosure_docx.py
"""
from __future__ import annotations

import shutil
import textwrap
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent
DOCX_OUT = ROOT / "docs" / "SkillSight_中国大陆专利技术交底书.docx"
FIGURE_DIR = ROOT / "docs" / "SkillSight_专利附图_黑白流程图"
FIGURE_ZIP_OUT = ROOT / "docs" / "SkillSight_专利附图_黑白流程图.zip"


def set_cell_shading(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    set_run_east_asia_font(run, "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell)


def set_run_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    set_run_east_asia_font(run, "黑体")


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str = "", bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21) if text else None
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    set_run_east_asia_font(run, "宋体")


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_east_asia_font(run, "宋体")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_east_asia_font(run, "宋体")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_flow(doc: Document, title: str, nodes: list[str]) -> None:
    add_p(doc, title, bold=True)
    table = doc.add_table(rows=1, cols=(len(nodes) * 2 - 1))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, node in enumerate(nodes):
        set_cell_text(table.rows[0].cells[idx * 2], node, bold=True)
        if idx < len(nodes) - 1:
            set_cell_text(table.rows[0].cells[idx * 2 + 1], "→")
    doc.add_paragraph()


def add_vertical_flow(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    add_p(doc, title, bold=True)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "步骤/模块", bold=True)
    set_cell_text(table.rows[0].cells[1], "处理内容", bold=True)
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, bold=True)
        set_cell_text(cells[1], right)
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    add_title(doc, "SkillSight 中国大陆发明专利技术交底书")
    add_p(doc, "文件性质：发明专利技术交底材料")
    add_p(doc, "拟定名称：一种基于可追溯证据的学生技能画像生成、岗位就绪度评估及可验证技能声明方法、系统、设备和存储介质")
    add_p(doc, "技术领域：教育信息化、人工智能评估、自然语言处理、向量检索、职业能力画像、数据安全与访问控制。")
    add_p(doc, "版本：v1.0（根据 skill_sight.zip、现有代码库和公开检索结果整理）")
    doc.add_page_break()


def add_background(doc: Document) -> None:
    add_h(doc, "一、发明名称", 1)
    add_p(doc, "一种基于可追溯证据的学生技能画像生成、岗位就绪度评估及可验证技能声明方法、系统、设备和存储介质。")

    add_h(doc, "二、技术领域", 1)
    add_p(doc, "本发明涉及人工智能与教育信息化领域，尤其涉及一种面向高校学生职业发展场景的证据驱动技能评估、岗位需求对齐、学习路径生成、简历优化和第三方可验证技能声明的计算机实现方法及系统。")

    add_h(doc, "三、背景技术及现有问题", 1)
    add_p(doc, "高校学生在课程作业、项目代码、实习报告、简历、社团活动、竞赛作品等材料中产生大量能够反映能力的证据，但这些材料通常分散存储，缺少统一的结构化技能画像；用人单位看到的简历多为自述，难以验证其真实性和熟练度；学校也难以量化课程体系与就业市场需求之间的对应关系。")
    add_p(doc, "现有职业平台、人岗匹配系统、简历解析系统或劳动力市场看板通常存在以下不足：")
    add_bullets(doc, [
        "偏重关键词出现、标签相似度或候选人与岗位的匹配分数，难以区分“提到某技能”和“实际展示某技能”。",
        "推荐结论和评分过程缺乏可追溯证据链，无法回溯到原始学生作品中的具体片段、位置和完整性校验信息。",
        "当检索证据不足或大模型输出不可靠时，系统仍可能给出正向结论，存在幻觉和误判风险。",
        "学生技能差距诊断与实际补缺动作割裂，无法自动映射到本校课程、项目、评估任务、简历修改等后续行动。",
        "导出的技能报告通常是静态 PDF 或自述文本，第三方难以在不访问学生隐私证据的情况下验证该报告是否由可信系统签发。",
        "高校内部部署场景需要基于角色和属性的访问控制、同意管理、撤回后的级联删除及审计追踪，通用招聘系统通常未形成闭环。"
    ])


def add_prior_art(doc: Document) -> None:
    add_h(doc, "四、公开专利初步检索及差异化方向", 1)
    add_p(doc, "基于公开网络和 Google Patents 的初筛，相近方向主要集中在人岗匹配、简历解析、人才画像、职业规划和学习路径推荐。建议正式申请前由代理机构补充 CNIPA/商业数据库检索。")
    add_table(doc, ["序号", "公开文献", "公开内容概括", "与本方案的主要区别"], [
        ["1", "CN120430763A 基于人工智能的自动化职位匹配系统和方法", "使用大语言模型、向量数据库和评分模块解析简历与岗位描述，按职位要求给候选人实时评分。", "本方案不是仅对简历与岗位进行评分，而是从学生真实作品构建带 quote_hash 的证据指针，采用 fail-closed 检索与输出护栏，多源证据融合后生成技能画像，并可映射校内课程和签发可公开验证的技能声明。"],
        ["2", "CN119313304A 一种基于多维度分析的职位简历智能匹配方法及系统", "提取简历和招聘信息的硬性/软性维度，计算求职者适配度分、企业适配度分和匹配矩阵。", "本方案重点在“证据是否足以证明技能”及技能熟练度，而非双向适配分；每个技能结论均绑定原始证据片段、字符范围和哈希，证据不足时拒绝判断。"],
        ["3", "CN111737485A 基于知识图谱、深度学习的人岗匹配方法、人岗匹配系统", "将简历、岗位及交互记录等有效信息加入知识图谱，训练推理模型输出匹配分数。", "本方案不依赖已选简历与岗位关系训练最优推理模型，而是以学生作品证据、技能 rubric、检索阈值、LLM 结构化输出和服务端护栏形成可审计评估链。"],
        ["4", "CN112990887A 一种简历和岗位匹配的方法及计算设备", "融合标签相关性向量与人岗嵌入向量，计算简历文本和岗位文本的人岗匹配度。", "本方案除技能与岗位匹配外，还包括同意门控、证据完整性校验、熟练度聚合、冲突检测、学习路径和 HMAC 验签导出。"],
        ["5", "CN120218048A 一种简历数据解析方法、装置、电子设备以及存储介质", "按行业类别选择简历解析模型，结合辅助特征和知识图谱获得结构化简历数据。", "本方案的输入不局限于简历，覆盖课程作业、项目、代码、报告等多类学生作品；核心不是字段解析，而是从作品证据到技能证明的可追溯判断。"],
        ["6", "CN110443571A 基于知识图谱进行简历评估的方法、装置及设备", "利用历史简历知识图谱标注非结构化简历信息并评估岗位相关性。", "本方案增加证据指针、quote_hash、同意撤回级联删除、检索可靠性分层、输出护栏及第三方验签导出等部署型可信机制。"],
        ["7", "CN116596412A 人才类型画像的实现方法及系统", "处理人才原始数据，构建人才知识图谱和能力模型，输出人才画像。", "本方案面向高校学生个人成长闭环，强调单个技能结论到具体证据片段的可验证性，并连接校内课程补缺、简历中心和职业顾问摘要。"],
        ["8", "CN118071032B 基于人工智能的个性化职业规划方法、装置及存储介质", "建立个人档案，多智能体分析职业要素、推荐动态职位并预测职业发展。", "本方案的创造性重点不在多智能体规划，而在证据驱动技能评估、检索失败拒绝、可信导出和高校治理架构。"],
    ])
    add_p(doc, "建议权利要求撰写时避免将发明概括为“AI 职位匹配”或“AI 简历评分”。更合适的差异化主线是：以同意授权的学生作品为输入，生成可完整性校验的证据指针；基于阈值检索、可选重排序和结构化拒绝实现 fail-closed；通过服务端护栏约束大模型输出；通过多源证据融合与冲突检测形成稳定技能画像；进一步将技能画像与岗位需求、校内课程、简历优化和可验证导出闭环连接。")


def add_solution(doc: Document) -> None:
    add_h(doc, "五、要解决的技术问题", 1)
    add_numbered(doc, [
        "如何将分散的学生学习成果和职业材料转化为可检索、可定位、可验证的技能证据库。",
        "如何在证据不足或证据质量不达标时阻止系统输出技能正向结论，降低大模型幻觉和错误匹配。",
        "如何区分技能“被提及”“被证明展示”和“熟练度达到特定水平”。",
        "如何将个人技能证据与岗位市场需求、校内课程资源和简历优化动作连接为可执行的发展路径。",
        "如何在高校部署环境下实现学生同意、数据撤回、权限隔离、审计和第三方最小披露验证。"
    ])

    add_h(doc, "六、技术方案概要", 1)
    add_p(doc, "本发明提供一种计算机实现的方法和系统。系统接收经学生授权的多模态作品或职业材料，将材料解析为带位置元数据和哈希值的证据块；基于技能定义或岗位需求生成检索请求，在向量数据库中检索相关证据块，并通过阈值过滤、可选重排序和可靠性计算决定是否进入大模型评估；大模型输出的技能展示标签、熟练度等级和证据引用必须通过服务端护栏验证，未通过则拒绝落库；系统对来自文档、交互评估、简历验证等不同来源的证据进行融合、时间衰减、权重计算和冲突检测，生成稳定技能画像；再将技能画像与岗位要求、课程技能映射和简历优化模块连接，生成岗位就绪度、课程补缺、学习路径、行动卡和简历改进建议；最后通过 HMAC 等密码学签名生成可由第三方公开验证的技能声明令牌，而不暴露原始学生证据。")

    add_flow(doc, "图1 系统总体架构黑白框图", [
        "学生/教师/管理端",
        "BFF 接口层",
        "同意与权限控制",
        "解析/嵌入 Worker",
        "证据库/向量库",
        "技能评估与护栏",
        "岗位/课程/简历/导出服务",
    ])
    add_vertical_flow(doc, "图2 六层技术流程图", [
        ("S1 证据上传", "接收文档、代码、简历、报告、图片、音视频等材料，并记录 purpose/scope 同意。"),
        ("S2 解析与证据映射", "生成 chunk_id、doc_id、字符偏移、页码、section_path、snippet 和 quote_hash。"),
        ("S3 市场与课程映射", "维护岗位技能需求、课程-技能映射、职位市场技能需求统计。"),
        ("S4 技能评估与验证", "基于检索证据、大模型结构化输出、服务端护栏和 rubric 生成技能展示与熟练度。"),
        ("S5 对话补证", "当静态证据不足时通过受控对话追问，形成带来源的补充证据。"),
        ("S6 推荐与支持", "输出岗位就绪度、课程补缺、学习路径、简历优化、顾问摘要和可验证导出。"),
    ])


def add_core_tech(doc: Document) -> None:
    add_h(doc, "七、核心技术组成", 1)
    add_h(doc, "7.1 同意门控的证据上传与级联删除", 2)
    add_p(doc, "每份材料上传时绑定用户、用途 purpose、范围 scope、授权状态和审计 request_id。后续解析、嵌入、检索、评估和展示均先查询授权状态；若授权不存在、过期或撤回，系统返回结构化拒绝，不进入下游处理。撤回同意后，系统删除向量、原始文件引用、证据块及与该文档绑定的评估记录。")
    add_flow(doc, "图3 同意门控与撤回级联删除图", [
        "授权 granted",
        "材料上传",
        "分块/嵌入",
        "检索/评估",
        "撤回 consent",
        "删除向量/文件/派生数据",
    ])

    add_h(doc, "7.2 证据指针与完整性校验", 2)
    add_p(doc, "系统将原始材料解析为证据块，每个证据块具有稳定标识和位置元数据。技能结论必须包含至少一个 EvidencePointer，或明确拒绝判断。quote_hash 为原始 chunk_text 的 SHA-256 值，用于防止展示片段被篡改或与原文不一致。")
    add_table(doc, ["字段", "含义"], [
        ["doc_id", "源文档标识"],
        ["chunk_id", "证据块标识"],
        ["char_start / char_end", "原文字符范围"],
        ["page_start / page_end", "页码范围，可选"],
        ["section_path", "文档章节路径，可选"],
        ["snippet", "不超过预设长度的展示片段"],
        ["quote_hash", "原始证据块文本哈希"],
        ["created_at", "指针创建时间"],
    ])
    add_flow(doc, "图4 证据指针生成与校验图", [
        "原始材料",
        "解析分块",
        "计算 quote_hash",
        "生成 EvidencePointer",
        "评估引用",
        "第三方/人工校验",
    ])

    add_h(doc, "7.3 Fail-closed 检索与可靠性分层", 2)
    add_p(doc, "系统根据技能定义、岗位要求或用户问题生成查询向量，从向量库取 top-k 证据块；先以最小相似度进行前置过滤，再可选调用重排序模型，之后进行后置阈值过滤。若剩余证据为空或分数边际不足，则输出包含 code、message、next_step 的 refusal，而不是将低质量证据交给大模型。")
    add_vertical_flow(doc, "图5 Fail-closed 检索决策图", [
        ("查询向量", "由技能定义、rubric 或岗位要求生成。"),
        ("Top-K 检索", "在向量库中获取候选证据块。"),
        ("前置阈值", "低于 similarity_min 的候选被过滤。"),
        ("可选重排序", "对候选证据重新排序并计算稳定性。"),
        ("后置阈值", "再次过滤不可靠证据。"),
        ("可靠性分层", "根据分数、边际、证据数量输出 high/medium/low。"),
        ("输出", "证据足够则进入评估；否则返回结构化拒绝。"),
    ])

    add_h(doc, "7.4 大模型输出护栏与技能判定", 2)
    add_p(doc, "大模型仅作为受约束的判断组件，不能直接决定落库结果。系统要求其输出固定 JSON schema，包括 demonstrated、mentioned、not_enough_information 等标签、evidence_chunk_ids、rationale、refusal_reason、matched_criteria 等字段。服务端验证 evidence_chunk_ids 必须属于检索允许集合；拒绝标签必须为空证据；正向标签必须有证据；熟练度 0 不得包含伪造 criteria。")
    add_flow(doc, "图6 大模型评估与服务端护栏图", [
        "允许证据集合",
        "LLM JSON 输出",
        "Schema 校验",
        "证据 ID 白名单校验",
        "标签/证据逻辑校验",
        "落库或拒绝",
    ])

    add_h(doc, "7.5 多源证据融合、冲突检测与稳定技能画像", 2)
    add_p(doc, "系统将文档证据、交互式评估、简历验证结果和人工复核结果作为不同来源。每条证据带有来源类型、时间戳、技能 ID、等级、置信度和证据指针。聚合器按来源权重和时间衰减计算候选等级，检查最低证据数量、跨证据一致性和正反证据冲突；若冲突超过阈值，则标记 needs_human_review 或降低可靠性。")
    add_vertical_flow(doc, "图7 多源技能等级聚合图", [
        ("文档评估", "来自学生作品的技能展示和熟练度证据。"),
        ("交互评估", "来自受控问答、编程、写作等测评结果。"),
        ("简历验证", "对简历声明进行证据检索和覆盖度判定。"),
        ("权重计算", "来源权重、时间衰减、可靠性权重。"),
        ("冲突检测", "等级分歧、互斥证据、证据数量不足。"),
        ("画像输出", "技能等级、可靠性、变更日志、人工复核标记。"),
    ])

    add_h(doc, "7.6 岗位需求、课程技能映射与学习路径", 2)
    add_p(doc, "系统维护角色库和 role_skill_requirements，每个岗位要求由 skill_id、目标等级和权重表示；同时维护 course_skill_map，将学校课程映射到技能。系统将用户当前技能等级与目标岗位要求左连接，计算 gap，并按重要性、缺口大小、市场需求和课程可获得性排序，输出课程补缺、项目建议、评估任务和阶段性里程碑。")
    add_flow(doc, "图8 岗位就绪度与学习路径生成图", [
        "目标岗位",
        "岗位技能要求",
        "用户技能画像",
        "缺口计算",
        "课程/项目/评估映射",
        "学习路径与行动卡",
    ])

    add_h(doc, "7.7 简历优化中心与声明验证", 2)
    add_p(doc, "简历优化中心并非仅生成文本建议，而是先解析简历声明，再对每条声明调用证据检索与验证服务，判断 supported、partial 或 unsupported，同时检查时间线、指标和经历一致性；随后按影响力、相关性、结构、语言、技能呈现和 ATS 可读性等 rubric 打分，生成改写建议、模板化输出和前后对比。")
    add_flow(doc, "图9 简历声明验证与优化流程图", [
        "简历上传",
        "声明抽取",
        "逐声明证据检索",
        "支持度判定",
        "Rubric 评分",
        "建议/模板/对比",
    ])

    add_h(doc, "7.8 HMAC 可验证技能声明", 2)
    add_p(doc, "用户可导出技能声明。系统对声明摘要、技能列表哈希、生成时间、有效期和发行方信息构成 payload，并使用服务端密钥进行 HMAC-SHA256 签名生成 token。第三方通过公开验证接口提交 token，系统只返回有效性、生成窗口和最小元数据，不返回学生原始证据。")
    add_flow(doc, "图10 可验证导出与公开验签图", [
        "技能声明摘要",
        "skills_hash",
        "HMAC 签名",
        "导出 token",
        "第三方 verify",
        "有效/过期/无效",
    ])

    add_h(doc, "7.9 RBAC + ABAC 访问控制与审计", 2)
    add_p(doc, "系统基于 student、staff、programme_leader、admin、career_coach 等角色配置动作权限，同时根据 programme_id、course_ids、term_id、teaching_relation 等属性判断访问范围。聚合接口对原始 chunk text、snippet、storage path、embedding、直接学生标识等字段进行 denylist 脱敏；审计中间件记录请求、主体、对象、状态和错误。")
    add_flow(doc, "图11 权限决策与脱敏输出图", [
        "JWT/身份",
        "角色权限",
        "属性范围",
        "Purpose 校验",
        "字段脱敏",
        "审计日志",
    ])


def add_effects_and_claims(doc: Document) -> None:
    add_h(doc, "八、有益效果", 1)
    add_bullets(doc, [
        "提高技能评估可信度：每个技能结论均可追溯到具体证据块、位置和完整性哈希。",
        "降低大模型误判：检索证据不足或输出不合规时 fail-closed 拒绝，而非继续生成正向结论。",
        "提高岗位建议的可解释性：岗位缺口来自用户真实证据与岗位技能要求之间的显式比较。",
        "形成学习闭环：从证据诊断、课程补缺、项目行动、交互测评、简历优化到再次验证。",
        "保护隐私和满足高校治理：通过同意门控、撤回级联删除、RBAC/ABAC、字段脱敏和审计日志降低数据风险。",
        "支持第三方低披露验证：用 HMAC token 验证技能声明真实性，无需暴露学生原始作品和敏感证据。"
    ])

    add_h(doc, "九、建议重点保护的创新点", 1)
    add_numbered(doc, [
        "同意门控的学生作品证据处理方法：授权状态作为解析、嵌入、检索、评估和展示的前置条件，撤回后级联删除派生数据。",
        "带完整性哈希的 EvidencePointer 数据结构及其在技能评估中的强制引用机制。",
        "面向技能证据检索的 fail-closed 管线：前置阈值、可选重排序、后置阈值、可靠性分层和结构化拒绝。",
        "大模型技能评估输出的服务端护栏方法：证据 ID 白名单、标签-证据一致性、rubric criteria 合规性验证。",
        "多源技能等级聚合方法：来源权重、时间衰减、最低证据数量、一致性检查、冲突检测和人工复核标记。",
        "基于可验证技能画像的岗位就绪度和课程补缺方法，将 role_skill_requirements、course_skill_map 与学生技能画像统一到同一 skill_id 体系。",
        "简历声明逐条证据验证与简历优化闭环，包括声明抽取、证据支持度判定、rubric 评分、建议生成和前后对比。",
        "技能声明的 HMAC 签名导出与公开验签方法，在不披露原始证据的前提下验证声明由系统签发。"
    ])

    add_h(doc, "十、专利保护范围初步布局", 1)
    add_table(doc, ["权利要求类型", "建议保护主题"], [
        ["独立方法权利要求 1", "从授权学生作品生成可追溯技能画像的方法，覆盖上传授权、分块哈希、证据检索、护栏评估、聚合画像、岗位/课程/导出的完整主流程。"],
        ["从属权利要求", "EvidencePointer 字段、quote_hash 校验、fail-closed 阈值、结构化 refusal、LLM 输出 schema、证据 ID 白名单、冲突检测、时间衰减、课程补缺、简历声明验证、HMAC 验签等。"],
        ["独立系统权利要求", "包括同意管理模块、证据解析模块、向量检索模块、技能评估护栏模块、聚合模块、推荐模块、导出验签模块、访问控制模块。"],
        ["设备权利要求", "电子设备，包括处理器和存储器，存储器存储指令，处理器执行上述方法。"],
        ["存储介质权利要求", "计算机可读存储介质，存储使计算设备执行上述方法的程序。"],
    ])


def add_implementation(doc: Document) -> None:
    add_h(doc, "十一、实施例", 1)
    add_h(doc, "实施例一：学生项目报告的技能画像生成", 2)
    add_numbered(doc, [
        "学生选择用途为“技能评估/岗位对齐”，上传项目报告和代码文件。",
        "系统解析文件，生成 chunk，并为每个 chunk 记录 doc_id、chunk_id、字符偏移、snippet 和 quote_hash。",
        "系统针对“Python 数据分析”技能生成查询，检索相关代码和报告段落。",
        "若检索得分满足阈值，大模型判断该技能为 demonstrated，并给出证据块 ID；服务端验证 ID 均来自允许集合后落库。",
        "熟练度评估根据 rubric 输出 level 2 或 level 3，聚合器结合历史证据输出稳定技能等级和可靠性。",
        "系统将该技能与目标岗位要求比较，若岗位要求 level 3 而学生为 level 2，则生成强化建议和课程推荐。"
    ])
    add_h(doc, "实施例二：简历声明验证和优化", 2)
    add_numbered(doc, [
        "学生上传简历，系统抽取“熟练掌握 Python 并完成数据可视化项目”等声明。",
        "系统对每条声明调用检索管线，查找学生已授权作品中的支持证据。",
        "若证据块包含相关代码和项目结果，声明标记为 supported；若仅出现课程名称或关键词，标记为 partial 或 unsupported。",
        "系统按简历 rubric 给出分项评分和修改建议，并将未被证据支持的夸大表述改为可验证表述。",
        "学生应用模板后再次评分，系统输出前后对比。"
    ])
    add_h(doc, "实施例三：第三方验证技能声明", 2)
    add_numbered(doc, [
        "学生导出技能声明，系统生成技能摘要并计算 skills_hash。",
        "系统使用服务端密钥对 payload 进行 HMAC 签名，生成带有效期的 token。",
        "用人单位访问公开 verify 接口提交 token。",
        "系统验证签名、时间窗和 payload 完整性，返回 valid、issued_at、expires_at 等最小元数据，不返回原始证据。"
    ])

    add_h(doc, "十二、可替代实施方式", 1)
    add_bullets(doc, [
        "向量数据库可为 Qdrant、pgvector、Milvus、FAISS 或其他近似最近邻索引。",
        "嵌入模型可为本地模型、云端模型或多语言模型；重排序器可为 cross-encoder、LLM reranker 或规则评分器。",
        "哈希算法可采用 SHA-256、SHA-3 或其他抗碰撞哈希；签名可采用 HMAC、非对称签名或机构证书。",
        "证据来源可扩展到 Git 提交、学习管理系统作业、在线评测结果、实习评价、竞赛证书和导师确认记录。",
        "岗位市场数据可来自公开招聘网站、校友就业数据、企业岗位库或第三方劳动力市场 API。",
        "权限模型可根据学校组织结构扩展学院、课程、项目、导师、职业顾问、企业合作伙伴等属性。"
    ])

    add_h(doc, "十三、术语表", 1)
    add_table(doc, ["术语", "解释"], [
        ["证据块 chunk", "由原始材料解析出的可定位文本或多模态转写片段。"],
        ["EvidencePointer", "指向证据块及其位置、摘要和哈希的结构化指针。"],
        ["fail-closed", "在证据不足或输出不合规时拒绝判断，而不是默认通过。"],
        ["rubric", "技能熟练度评估规则，通常包括等级和 criteria。"],
        ["岗位就绪度", "用户当前技能画像与目标岗位技能要求之间的满足程度。"],
        ["HMAC token", "由服务端密钥签名、可公开验证的声明令牌。"],
    ])


def wrap_svg_text(text: str, width: int = 10) -> list[str]:
    lines: list[str] = []
    for part in text.split("\n"):
        lines.extend(textwrap.wrap(part, width=width, break_long_words=False) or [""])
    return lines[:4]


def svg_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def svg_box(x: int, y: int, w: int, h: int, label: str) -> str:
    lines = wrap_svg_text(label, width=max(4, w // 18))
    start_y = y + h / 2 - (len(lines) - 1) * 9
    text = "\n".join(
        f'<text x="{x + w / 2}" y="{start_y + idx * 18:.0f}" text-anchor="middle" '
        f'font-size="14" font-family="Arial, SimSun, sans-serif">{svg_escape(line)}</text>'
        for idx, line in enumerate(lines)
    )
    return (
        f'<rect x="{x}" y="{y}" width="{w}" height="{h}" fill="white" '
        f'stroke="black" stroke-width="1.6"/>\n{text}'
    )


def svg_diamond(cx: int, cy: int, w: int, h: int, label: str) -> str:
    points = f"{cx},{cy - h // 2} {cx + w // 2},{cy} {cx},{cy + h // 2} {cx - w // 2},{cy}"
    lines = wrap_svg_text(label, width=max(4, w // 20))
    start_y = cy - (len(lines) - 1) * 8
    text = "\n".join(
        f'<text x="{cx}" y="{start_y + idx * 16:.0f}" text-anchor="middle" '
        f'font-size="13" font-family="Arial, SimSun, sans-serif">{svg_escape(line)}</text>'
        for idx, line in enumerate(lines)
    )
    return f'<polygon points="{points}" fill="white" stroke="black" stroke-width="1.6"/>\n{text}'


def svg_arrow(x1: int, y1: int, x2: int, y2: int, label: str | None = None) -> str:
    mid = ""
    if label:
        mid = (
            f'<text x="{(x1 + x2) / 2}" y="{(y1 + y2) / 2 - 6}" text-anchor="middle" '
            f'font-size="12" font-family="Arial, SimSun, sans-serif">{svg_escape(label)}</text>'
        )
    return (
        f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="black" '
        f'stroke-width="1.4" marker-end="url(#arrow)"/>\n{mid}'
    )


def render_horizontal_svg(title: str, nodes: list[str], filename: str) -> None:
    box_w, box_h, gap = 130, 70, 42
    width = 80 + len(nodes) * box_w + (len(nodes) - 1) * gap
    height = 180
    y = 70
    parts = [svg_header(width, height, title)]
    for idx, node in enumerate(nodes):
        x = 40 + idx * (box_w + gap)
        parts.append(svg_box(x, y, box_w, box_h, node))
        if idx < len(nodes) - 1:
            parts.append(svg_arrow(x + box_w, y + box_h // 2, x + box_w + gap, y + box_h // 2))
    parts.append("</svg>")
    (FIGURE_DIR / filename).write_text("\n".join(parts), encoding="utf-8")


def render_vertical_svg(title: str, rows: list[str], filename: str) -> None:
    box_w, box_h, gap = 420, 54, 26
    width = 560
    height = 90 + len(rows) * box_h + (len(rows) - 1) * gap + 40
    x = 70
    parts = [svg_header(width, height, title)]
    for idx, row in enumerate(rows):
        y = 70 + idx * (box_h + gap)
        parts.append(svg_box(x, y, box_w, box_h, row))
        if idx < len(rows) - 1:
            parts.append(svg_arrow(x + box_w // 2, y + box_h, x + box_w // 2, y + box_h + gap))
    parts.append("</svg>")
    (FIGURE_DIR / filename).write_text("\n".join(parts), encoding="utf-8")


def render_decision_svg(title: str, nodes: list[str], decision: str, filename: str) -> None:
    width, height = 940, 360
    parts = [svg_header(width, height, title)]
    x0, y = 60, 110
    box_w, box_h, gap = 150, 62, 36
    for idx, node in enumerate(nodes[:3]):
        x = x0 + idx * (box_w + gap)
        parts.append(svg_box(x, y, box_w, box_h, node))
        parts.append(svg_arrow(x + box_w, y + box_h // 2, x + box_w + gap, y + box_h // 2))
    d_cx = x0 + 3 * (box_w + gap) + 55
    parts.append(svg_diamond(d_cx, y + box_h // 2, 120, 92, decision))
    parts.append(svg_arrow(d_cx + 60, y + box_h // 2, d_cx + 125, y + box_h // 2, "是"))
    parts.append(svg_box(d_cx + 125, y, 160, box_h, "进入下游处理"))
    parts.append(svg_arrow(d_cx, y + 46, d_cx, y + 120, "否"))
    parts.append(svg_box(d_cx - 80, y + 120, 160, box_h, "结构化拒绝\ncode/message/next_step"))
    parts.append("</svg>")
    (FIGURE_DIR / filename).write_text("\n".join(parts), encoding="utf-8")


def svg_header(width: int, height: int, title: str) -> str:
    return f'''<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">
<defs>
  <marker id="arrow" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth">
    <path d="M0,0 L0,6 L9,3 z" fill="black"/>
  </marker>
</defs>
<rect x="1" y="1" width="{width - 2}" height="{height - 2}" fill="white" stroke="black" stroke-width="1"/>
<text x="{width / 2}" y="36" text-anchor="middle" font-size="20" font-weight="bold" font-family="Arial, SimHei, sans-serif">{svg_escape(title)}</text>'''


def build_figure_zip() -> None:
    if FIGURE_DIR.exists():
        shutil.rmtree(FIGURE_DIR)
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    render_horizontal_svg("图1 系统总体架构图", [
        "用户端", "BFF 接口层", "同意与权限控制", "解析/嵌入 Worker", "证据库/向量库", "评估护栏", "推荐与导出服务"
    ], "图1_系统总体架构图.svg")
    render_vertical_svg("图2 六层技术流程图", [
        "S1 证据上传与授权", "S2 解析分块与证据映射", "S3 岗位市场与课程映射", "S4 技能评估与验证", "S5 对话补证", "S6 推荐、简历与可验证导出"
    ], "图2_六层技术流程图.svg")
    render_horizontal_svg("图3 同意门控与级联删除图", [
        "授权 granted", "材料上传", "分块/嵌入", "检索/评估", "撤回同意", "级联删除派生数据"
    ], "图3_同意门控与级联删除图.svg")
    render_horizontal_svg("图4 证据指针生成与完整性校验图", [
        "原始材料", "解析分块", "计算 quote_hash", "生成 EvidencePointer", "评估引用", "完整性校验"
    ], "图4_证据指针生成与完整性校验图.svg")
    render_decision_svg("图5 Fail-closed 检索决策图", [
        "查询向量", "Top-K 检索", "阈值/重排序"
    ], "证据是否充分", "图5_Fail_closed检索决策图.svg")
    render_decision_svg("图6 大模型输出护栏图", [
        "允许证据集合", "LLM JSON 输出", "Schema/白名单校验"
    ], "逻辑是否合规", "图6_大模型输出护栏图.svg")
    render_vertical_svg("图7 多源证据融合与冲突检测图", [
        "文档评估证据", "交互评估证据", "简历验证证据", "来源权重与时间衰减", "一致性与冲突检测", "稳定技能画像"
    ], "图7_多源证据融合与冲突检测图.svg")
    render_horizontal_svg("图8 岗位就绪度与学习路径生成图", [
        "目标岗位", "岗位技能要求", "用户技能画像", "缺口计算", "课程/项目/评估映射", "学习路径"
    ], "图8_岗位就绪度与学习路径生成图.svg")
    render_horizontal_svg("图9 简历声明验证与优化流程图", [
        "简历上传", "声明抽取", "逐声明证据检索", "支持度判定", "Rubric 评分", "建议/模板/对比"
    ], "图9_简历声明验证与优化流程图.svg")
    render_horizontal_svg("图10 可验证导出与公开验签图", [
        "技能声明摘要", "skills_hash", "HMAC 签名", "导出 token", "公开 verify", "有效/过期/无效"
    ], "图10_可验证导出与公开验签图.svg")
    render_horizontal_svg("图11 RBAC+ABAC 权限决策图", [
        "JWT/身份", "角色权限", "属性范围", "Purpose 校验", "字段脱敏", "审计日志"
    ], "图11_RBAC_ABAC权限决策图.svg")
    if FIGURE_ZIP_OUT.exists():
        FIGURE_ZIP_OUT.unlink()
    with zipfile.ZipFile(FIGURE_ZIP_OUT, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(FIGURE_DIR.glob("*.svg")):
            archive.write(path, arcname=path.name)
    shutil.rmtree(FIGURE_DIR)


def build_docx() -> None:
    doc = Document()
    style_doc(doc)
    add_cover(doc)
    add_background(doc)
    add_prior_art(doc)
    add_solution(doc)
    add_core_tech(doc)
    add_effects_and_claims(doc)
    add_implementation(doc)
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    build_figure_zip()
    print(f"Saved: {DOCX_OUT}")
    print(f"Saved: {FIGURE_ZIP_OUT}")


if __name__ == "__main__":
    build_docx()
